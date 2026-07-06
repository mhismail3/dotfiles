#!/usr/bin/env python3
"""Private Personal OS helper for Codex.

This script is intentionally file-first and reversible. It writes only inside
~/.codex/personal-os unless a user explicitly invokes Things writeback or
LaunchAgent automation commands.
"""

from __future__ import annotations

import argparse
import datetime as dt
import fnmatch
import hashlib
import html.parser
import json
import mimetypes
import os
import plistlib
import re
import secrets
import shutil
import subprocess
import sys
import tempfile
from pathlib import Path
from typing import Any

DEFAULT_ROOT = Path("~/.codex/personal-os").expanduser()
CHRONICLE_RESOURCES = Path("~/.codex/memories/extensions/chronicle/resources").expanduser()
THINGS_SCRIPT = Path("~/.codex/skills/things-3/scripts/things.py").expanduser()
CALENDAR_SCRIPT = Path("~/.codex/skills/apple-calendar/scripts/calendar.py").expanduser()
LLM_WIKI_ROOT = Path("~/.codex/wiki").expanduser()
LAUNCH_AGENT_LABEL = "com.moose.personal-os.daily"
LAUNCH_AGENT_PATH = Path("~/Library/LaunchAgents").expanduser() / f"{LAUNCH_AGENT_LABEL}.plist"
BUNDLED_PYTHON = Path("~/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3").expanduser()
BUNDLED_BIN = BUNDLED_PYTHON.parent.parent.parent / "bin"
PDFTOPPM = BUNDLED_BIN / "pdftoppm"
THINGS_PROJECT_NAME = "🤖 Agent Tasks"
THINGS_PROJECT_ID = "Du31k5uGbRWNgX6F4WFAnK"
THINGS_TAG = "Personal OS"
LARGE_FILE_THRESHOLD_BYTES = 250 * 1024 * 1024
THINGS_TIMEOUT_SECONDS = int(os.environ.get("PERSONAL_OS_THINGS_TIMEOUT", "15"))

REQUIRED_DIRS = (
    "inbox",
    "journal",
    "episodes",
    "reflections/daily",
    "profile",
    "files",
    "files/originals",
    "files/manifests",
    "files/extracted",
    "files/summaries",
    "files/outlines",
    "files/indexes",
    "_views",
    "_state",
    "_state/launchagents",
    "_state/time",
    "_reports",
    "_reports/garden",
    "_reports/time",
    "_logs",
    "_logs/runs",
)

FILE_KINDS = (
    "household",
    "financial",
    "medical",
    "legal",
    "identity",
    "career",
    "home",
    "travel",
    "procedure",
    "reference",
    "archive",
    "other",
)

SENSITIVE_FILE_KINDS = {"financial", "medical", "legal", "identity"}
SENSITIVITY_LABELS = ("normal", "sensitive")
EXTRACTION_STATUSES = {"pending", "complete", "failed", "needs_ocr", "unsupported"}
SUMMARY_STATUSES = {"pending", "complete", "failed", "pending_agent"}
BULK_EXCLUDE_PATTERNS = (
    ".git",
    "node_modules",
    "*.app",
    "*.photoslibrary",
    "*.sparsebundle",
    "*.backupbundle",
    "Library",
    "Applications",
)
TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".markdown",
    ".csv",
    ".tsv",
    ".json",
    ".jsonl",
    ".yaml",
    ".yml",
    ".xml",
    ".log",
}
HTML_EXTENSIONS = {".html", ".htm"}
RTF_EXTENSIONS = {".rtf"}
DOCX_EXTENSIONS = {".docx"}
LEGACY_WORD_EXTENSIONS = {".doc"}
XLSX_EXTENSIONS = {".xlsx"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".tif", ".tiff", ".heic", ".webp"}

PROFILE_FILES = {
    "preference": "preferences.md",
    "pattern": "patterns.md",
    "routine": "routines.md",
    "value": "values.md",
}

ACTION_RE = re.compile(
    r"^\s*(?:[-*]\s*)?(?:\[\s\]\s*)?(?:(ACTION|TODO|NEXT|FOLLOW\s+UP)\s*[:\-]\s*)(.+?)\s*$",
    re.I,
)
CHECKBOX_RE = re.compile(r"^\s*[-*]\s+\[\s\]\s+(.+?)\s*$")
PROFILE_RE = re.compile(r"^\s*(?:[-*]\s*)?(preference|pattern|routine|value)\s*[:\-]\s*(.+?)\s*$", re.I)
UTC = dt.timezone.utc


def utc_now() -> dt.datetime:
    return dt.datetime.now(UTC).replace(microsecond=0)


def today() -> str:
    return dt.date.today().isoformat()


def parse_day(value: str | None) -> dt.date:
    if not value or value == "today":
        return dt.date.today()
    if value == "yesterday":
        return dt.date.today() - dt.timedelta(days=1)
    try:
        return dt.date.fromisoformat(value)
    except ValueError as exc:
        raise SystemExit(f"Use YYYY-MM-DD, today, or yesterday for dates: {value}") from exc


def slugify(value: str, fallback: str = "entry") -> str:
    value = value.casefold()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    value = re.sub(r"-+", "-", value).strip("-")
    return value[:80] or fallback


def json_quote(value: str) -> str:
    return json.dumps(value, ensure_ascii=False)


def frontmatter(fields: dict[str, Any]) -> str:
    lines = ["---"]
    for key, value in fields.items():
        if isinstance(value, list):
            rendered = "[" + ", ".join(json_quote(str(item)) for item in value) + "]"
        elif value is None:
            rendered = '""'
        elif isinstance(value, bool):
            rendered = "true" if value else "false"
        else:
            rendered = json_quote(str(value))
        lines.append(f"{key}: {rendered}")
    lines.append("---")
    return "\n".join(lines) + "\n"


def print_json(data: Any) -> None:
    print(json.dumps(data, indent=2, ensure_ascii=False, sort_keys=True))


def root_from_args(args: argparse.Namespace) -> Path:
    return Path(args.root).expanduser()


def relative(root: Path, path: Path) -> str:
    try:
        return str(path.resolve().relative_to(root.resolve()))
    except ValueError:
        return str(path.expanduser())


def ensure_root_exists(root: Path) -> None:
    if not root.exists():
        raise SystemExit(f"Personal OS root does not exist. Run bootstrap first: {root}")


def ensure_under_root(root: Path, path: Path) -> Path:
    resolved = path.expanduser().resolve()
    try:
        resolved.relative_to(root.resolve())
    except ValueError as exc:
        raise SystemExit(f"Refusing path outside Personal OS root: {path}") from exc
    return resolved


def read_stdin(args: argparse.Namespace) -> str:
    if not getattr(args, "stdin", False):
        raise SystemExit("This command requires --stdin for input text.")
    return sys.stdin.read().rstrip() + "\n"


def run_id() -> str:
    return utc_now().strftime("%Y%m%dT%H%M%SZ") + "-" + secrets.token_hex(3)


def run_command(cmd: list[str], input_text: str | None = None, timeout: int | None = None) -> subprocess.CompletedProcess[str]:
    try:
        return subprocess.run(cmd, input=input_text, text=True, capture_output=True, check=False, timeout=timeout)
    except subprocess.TimeoutExpired as exc:
        stdout = exc.stdout if isinstance(exc.stdout, str) else (exc.stdout or b"").decode(errors="replace")
        stderr = exc.stderr if isinstance(exc.stderr, str) else (exc.stderr or b"").decode(errors="replace")
        message = stderr.strip() or stdout.strip() or f"Timed out after {timeout} seconds: {' '.join(cmd)}"
        return subprocess.CompletedProcess(cmd, 124, stdout or "", message)


def git_head(root: Path) -> str | None:
    if not (root / ".git").exists():
        return None
    proc = run_command(["git", "-C", str(root), "rev-parse", "--verify", "HEAD"])
    if proc.returncode != 0:
        return None
    return proc.stdout.strip()


def git_dirty(root: Path) -> bool:
    if not (root / ".git").exists():
        return False
    proc = run_command(["git", "-C", str(root), "status", "--porcelain"])
    return bool(proc.stdout.strip())


def ensure_private_git(root: Path) -> None:
    if not (root / ".git").exists():
        proc = run_command(["git", "-C", str(root), "init"])
        if proc.returncode != 0:
            raise SystemExit(proc.stderr.strip() or "git init failed")
    if run_command(["git", "-C", str(root), "config", "user.name"]).returncode != 0:
        run_command(["git", "-C", str(root), "config", "user.name", "Personal OS"])
    if run_command(["git", "-C", str(root), "config", "user.email"]).returncode != 0:
        run_command(["git", "-C", str(root), "config", "user.email", "personal-os@localhost"])


def commit_all(root: Path, message: str, no_commit: bool = False) -> str | None:
    if no_commit or not (root / ".git").exists():
        return git_head(root)
    run_command(["git", "-C", str(root), "add", "-A"])
    if not git_dirty(root):
        return git_head(root)
    proc = run_command(["git", "-C", str(root), "commit", "-m", message])
    if proc.returncode != 0:
        raise SystemExit(proc.stderr.strip() or "git commit failed")
    return git_head(root)


def load_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    return json.loads(path.read_text(encoding="utf-8"))


def write_text(path: Path, text: str, touched: list[Path]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    touched.append(path)


def write_json(path: Path, data: Any, touched: list[Path]) -> None:
    write_text(path, json.dumps(data, indent=2, ensure_ascii=False, sort_keys=True) + "\n", touched)


def append_text(path: Path, text: str, touched: list[Path]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("a", encoding="utf-8") as handle:
        handle.write(text)
    touched.append(path)


def append_section_if_missing(path: Path, marker: str, section: str, touched: list[Path]) -> None:
    existing = path.read_text(encoding="utf-8") if path.exists() else ""
    if marker in existing:
        return
    write_text(path, existing.rstrip() + "\n\n" + section.strip() + "\n", touched)


def file_manifest_path(root: Path, file_id: str) -> Path:
    return root / "files" / "manifests" / f"{file_id}.json"


def file_extracted_path(root: Path, file_id: str) -> Path:
    return root / "files" / "extracted" / f"{file_id}.txt"


def file_summary_path(root: Path, file_id: str) -> Path:
    return root / "files" / "summaries" / f"{file_id}.md"


def file_outline_path(root: Path, file_id: str) -> Path:
    return root / "files" / "outlines" / f"{file_id}.md"


def things_routing_path(root: Path) -> Path:
    return root / "_state" / "things-routing.json"


def default_things_routing() -> dict[str, str]:
    return {
        "destination": "project",
        "project_name": THINGS_PROJECT_NAME,
        "project_id": THINGS_PROJECT_ID,
        "tag": THINGS_TAG,
    }


def load_things_routing(root: Path) -> dict[str, str]:
    data = load_json(things_routing_path(root), default_things_routing())
    routing = default_things_routing()
    if isinstance(data, dict):
        routing.update({key: str(value) for key, value in data.items() if value is not None})
    return routing


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def detect_mime(path: Path) -> str:
    guessed = mimetypes.guess_type(path.name)[0]
    if guessed:
        return guessed
    proc = run_command(["file", "--brief", "--mime-type", str(path)])
    if proc.returncode == 0 and proc.stdout.strip():
        return proc.stdout.strip()
    return "application/octet-stream"


def file_title_from_path(path: Path, title: str | None = None) -> str:
    return title.strip() if title and title.strip() else path.stem.replace("-", " ").replace("_", " ").strip() or path.name


def file_id_for(path: Path, title: str, sha256: str, added: dt.date | None = None) -> str:
    day = added or dt.date.today()
    return f"{day.isoformat().replace('-', '')}-{slugify(title, 'file')}-{sha256[:12]}"


def sensitivity_for_kind(kind: str) -> str:
    return "sensitive" if kind in SENSITIVE_FILE_KINDS else "normal"


def source_modified_at(path: Path) -> str | None:
    try:
        return dt.datetime.fromtimestamp(path.stat().st_mtime, UTC).isoformat().replace("+00:00", "Z")
    except FileNotFoundError:
        return None


def load_file_manifests(root: Path) -> list[dict[str, Any]]:
    manifests: list[dict[str, Any]] = []
    for path in sorted((root / "files" / "manifests").glob("*.json")):
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            continue
        if isinstance(data, dict):
            manifests.append(data)
    manifests.sort(key=lambda item: str(item.get("added_at", "")), reverse=True)
    return manifests


def find_manifest(root: Path, file_id: str) -> dict[str, Any]:
    path = file_manifest_path(root, file_id)
    if not path.exists():
        raise SystemExit(f"No Personal OS file manifest found for {file_id}")
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise SystemExit(f"Invalid manifest: {relative(root, path)}")
    return data


def duplicate_for_sha(root: Path, sha256: str, file_id: str) -> str | None:
    for manifest in load_file_manifests(root):
        if manifest.get("sha256") == sha256 and manifest.get("file_id") != file_id:
            return str(manifest.get("file_id"))
    return None


def manifest_source_path(manifest: dict[str, Any]) -> Path | None:
    copied = manifest.get("copied_path")
    source = manifest.get("source_path")
    if copied:
        return Path(str(copied)).expanduser()
    if source:
        return Path(str(source)).expanduser()
    return None


def relative_or_empty(root: Path, value: str | None) -> str:
    if not value:
        return ""
    return relative(root, Path(value))


class _HTMLTextExtractor(html.parser.HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_data(self, data: str) -> None:
        clean = data.strip()
        if clean:
            self.parts.append(clean)

    def text(self) -> str:
        return "\n".join(self.parts)


def strip_html(text: str) -> str:
    parser = _HTMLTextExtractor()
    parser.feed(text)
    return parser.text()


def bundled_python_available() -> bool:
    return BUNDLED_PYTHON.exists() and os.access(BUNDLED_PYTHON, os.X_OK)


VISION_OCR_SWIFT = r'''
import Foundation
import Vision
import AppKit

func loadCGImage(_ path: String) -> CGImage? {
    let url = URL(fileURLWithPath: path)
    guard let image = NSImage(contentsOf: url),
          let tiff = image.tiffRepresentation,
          let bitmap = NSBitmapImageRep(data: tiff) else {
        return nil
    }
    return bitmap.cgImage
}

var hadError = false
var emitted = false

for path in CommandLine.arguments.dropFirst() {
    guard let cgImage = loadCGImage(path) else {
        fputs("Could not load image: \(path)\n", stderr)
        hadError = true
        continue
    }
    let request = VNRecognizeTextRequest()
    request.recognitionLevel = .accurate
    request.usesLanguageCorrection = true
    let handler = VNImageRequestHandler(cgImage: cgImage, options: [:])
    do {
        try handler.perform([request])
    } catch {
        fputs("OCR failed for \(path): \(error)\n", stderr)
        hadError = true
        continue
    }
    let observations = request.results ?? []
    let lines = observations.compactMap { observation in
        observation.topCandidates(1).first?.string
    }
    if !lines.isEmpty {
        if emitted {
            print("")
            print("")
        }
        print("## OCR: \(URL(fileURLWithPath: path).lastPathComponent)")
        print("")
        print(lines.joined(separator: "\n"))
        emitted = true
    }
}

if hadError && !emitted {
    exit(2)
}
'''


def vision_ocr_available() -> bool:
    return shutil.which("swift") is not None


def ocr_images_with_vision(paths: list[Path], method: str = "vision-ocr") -> dict[str, str]:
    if not paths:
        return {"status": "needs_ocr", "method": method, "text": "", "error": "No rendered images to OCR"}
    if not vision_ocr_available():
        return {"status": "needs_ocr", "method": method, "text": "", "error": "macOS Vision OCR provider is not available"}
    proc = run_command(["swift", "-", *[str(path) for path in paths]], input_text=VISION_OCR_SWIFT)
    text = proc.stdout.strip()
    if proc.returncode != 0 and not text:
        return {"status": "needs_ocr", "method": method, "text": "", "error": proc.stderr.strip() or proc.stdout.strip()}
    if text:
        return {"status": "complete", "method": method, "text": text, "error": proc.stderr.strip()}
    return {"status": "needs_ocr", "method": method, "text": "", "error": proc.stderr.strip() or "OCR completed without recognized text"}


def ocr_pdf_with_vision(path: Path) -> dict[str, str]:
    if not PDFTOPPM.exists():
        return {"status": "needs_ocr", "method": "vision-ocr-pdf", "text": "", "error": f"pdftoppm missing: {PDFTOPPM}"}
    with tempfile.TemporaryDirectory(prefix="personal-os-ocr-") as tmp:
        prefix = Path(tmp) / "page"
        proc = run_command([str(PDFTOPPM), "-r", "200", "-png", str(path), str(prefix)])
        if proc.returncode != 0:
            return {"status": "needs_ocr", "method": "vision-ocr-pdf", "text": "", "error": proc.stderr.strip() or proc.stdout.strip()}
        pages = sorted(Path(tmp).glob("page-*.png"))
        return ocr_images_with_vision(pages, "vision-ocr-pdf")


def extract_pdf(path: Path) -> dict[str, str]:
    if not bundled_python_available():
        return {"status": "failed", "method": "pdf", "text": "", "error": f"Bundled Python missing: {BUNDLED_PYTHON}"}
    code = r'''
import json, sys
path = sys.argv[1]
text = ""
method = "pdfplumber"
error = ""
try:
    import pdfplumber
    with pdfplumber.open(path) as pdf:
        text = "\n\n".join((page.extract_text() or "") for page in pdf.pages)
except Exception as exc:
    error = str(exc)
    try:
        from pypdf import PdfReader
        method = "pypdf"
        reader = PdfReader(path)
        text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        error = ""
    except Exception as exc2:
        error = f"{error}; pypdf: {exc2}" if error else str(exc2)
status = "complete" if text.strip() else ("needs_ocr" if not error else "failed")
print(json.dumps({"status": status, "method": method, "text": text, "error": error}))
'''
    proc = run_command([str(BUNDLED_PYTHON), "-c", code, str(path)])
    if proc.returncode != 0:
        return {"status": "failed", "method": "pdf", "text": "", "error": proc.stderr.strip() or proc.stdout.strip()}
    try:
        data = json.loads(proc.stdout)
    except json.JSONDecodeError:
        return {"status": "failed", "method": "pdf", "text": "", "error": proc.stdout[:500]}
    result = {key: str(data.get(key, "")) for key in ("status", "method", "text", "error")}
    if result.get("status") == "complete" and result.get("text", "").count("(cid:") > 50:
        ocr_result = ocr_pdf_with_vision(path)
        if ocr_result.get("status") == "complete":
            ocr_result["error"] = (
                ocr_result.get("error", "")
                + "\n"
                + f"{result.get('method', 'pdf')} extraction produced CID font artifacts; used Vision OCR fallback"
            ).strip()
            return ocr_result
    if result.get("status") == "needs_ocr":
        ocr_result = ocr_pdf_with_vision(path)
        if ocr_result.get("status") == "complete":
            if result.get("error"):
                ocr_result["error"] = (ocr_result.get("error", "") + "\n" + result["error"]).strip()
            return ocr_result
    return result


def extract_word_with_textutil(path: Path, method: str = "textutil-word") -> dict[str, str]:
    proc = run_command(["textutil", "-convert", "txt", "-stdout", str(path)])
    if proc.returncode != 0:
        return {"status": "failed", "method": method, "text": "", "error": proc.stderr.strip() or proc.stdout.strip()}
    return {"status": "complete" if proc.stdout.strip() else "failed", "method": method, "text": proc.stdout, "error": ""}


def extract_docx(path: Path) -> dict[str, str]:
    if not bundled_python_available():
        return {"status": "failed", "method": "python-docx", "text": "", "error": f"Bundled Python missing: {BUNDLED_PYTHON}"}
    code = r'''
import json, sys
path = sys.argv[1]
parts = []
error = ""
try:
    import docx
    doc = docx.Document(path)
    parts.extend(p.text for p in doc.paragraphs if p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
except Exception as exc:
    error = str(exc)
text = "\n".join(parts)
print(json.dumps({"status": "complete" if text.strip() else "failed", "method": "python-docx", "text": text, "error": error}))
'''
    proc = run_command([str(BUNDLED_PYTHON), "-c", code, str(path)])
    if proc.returncode != 0:
        return {"status": "failed", "method": "python-docx", "text": "", "error": proc.stderr.strip() or proc.stdout.strip()}
    try:
        data = json.loads(proc.stdout)
    except json.JSONDecodeError:
        return {"status": "failed", "method": "python-docx", "text": "", "error": proc.stdout[:500]}
    result = {key: str(data.get(key, "")) for key in ("status", "method", "text", "error")}
    if result.get("status") != "complete":
        fallback = extract_word_with_textutil(path, "textutil-docx")
        if fallback.get("status") == "complete":
            return fallback
    return result


def extract_xlsx(path: Path) -> dict[str, str]:
    if not bundled_python_available():
        return {"status": "failed", "method": "openpyxl", "text": "", "error": f"Bundled Python missing: {BUNDLED_PYTHON}"}
    code = r'''
import json, sys
path = sys.argv[1]
parts = []
error = ""
method = "openpyxl"
try:
    import openpyxl
    workbook = openpyxl.load_workbook(path, read_only=True, data_only=True)
    for worksheet in workbook.worksheets:
        parts.append(f"# Sheet: {worksheet.title}")
        for row in worksheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            while values and values[-1] == "":
                values.pop()
            if any(value.strip() for value in values):
                parts.append("\t".join(values))
        parts.append("")
except Exception as exc:
    error = str(exc)
    try:
        import posixpath
        import re
        import zipfile
        import xml.etree.ElementTree as ET

        ns = {
            "m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main",
            "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
            "rel": "http://schemas.openxmlformats.org/package/2006/relationships",
        }

        def read_xml(zf, name):
            return ET.fromstring(zf.read(name))

        def cell_column(cell_ref):
            match = re.match(r"([A-Z]+)", cell_ref or "")
            if not match:
                return 0
            value = 0
            for char in match.group(1):
                value = value * 26 + (ord(char) - ord("A") + 1)
            return value

        def text_content(element):
            return "".join(element.itertext()) if element is not None else ""

        with zipfile.ZipFile(path) as zf:
            shared = []
            if "xl/sharedStrings.xml" in zf.namelist():
                root = read_xml(zf, "xl/sharedStrings.xml")
                for item in root.findall("m:si", ns):
                    shared.append(text_content(item))

            workbook_root = read_xml(zf, "xl/workbook.xml")
            rels_root = read_xml(zf, "xl/_rels/workbook.xml.rels")
            rels = {}
            for rel in rels_root.findall("rel:Relationship", ns):
                target = rel.attrib.get("Target", "")
                if not target.startswith("/"):
                    target = posixpath.normpath(posixpath.join("xl", target))
                else:
                    target = target.lstrip("/")
                rels[rel.attrib.get("Id", "")] = target

            for sheet in workbook_root.findall("m:sheets/m:sheet", ns):
                title = sheet.attrib.get("name", "Sheet")
                rid = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", "")
                sheet_path = rels.get(rid)
                if not sheet_path or sheet_path not in zf.namelist():
                    continue
                sheet_root = read_xml(zf, sheet_path)
                parts.append(f"# Sheet: {title}")
                for row in sheet_root.findall("m:sheetData/m:row", ns):
                    cells = []
                    for cell in row.findall("m:c", ns):
                        col = cell_column(cell.attrib.get("r", ""))
                        while len(cells) < max(col - 1, 0):
                            cells.append("")
                        cell_type = cell.attrib.get("t", "")
                        if cell_type == "s":
                            raw = text_content(cell.find("m:v", ns))
                            value = shared[int(raw)] if raw.isdigit() and int(raw) < len(shared) else raw
                        elif cell_type == "inlineStr":
                            value = text_content(cell.find("m:is", ns))
                        else:
                            value = text_content(cell.find("m:v", ns))
                        cells.append(value)
                    while cells and cells[-1] == "":
                        cells.pop()
                    if any(value.strip() for value in cells):
                        parts.append("\t".join(cells))
                parts.append("")
        if parts:
            method = "ooxml-xlsx"
            error = f"openpyxl failed: {error}; used raw OOXML fallback"
    except Exception as fallback_exc:
        error = f"{error}; OOXML fallback failed: {fallback_exc}"
text = "\n".join(parts).strip()
print(json.dumps({"status": "complete" if text else "failed", "method": method, "text": text, "error": error}))
'''
    proc = run_command([str(BUNDLED_PYTHON), "-c", code, str(path)])
    if proc.returncode != 0:
        return {"status": "failed", "method": "openpyxl", "text": "", "error": proc.stderr.strip() or proc.stdout.strip()}
    try:
        data = json.loads(proc.stdout)
    except json.JSONDecodeError:
        return {"status": "failed", "method": "openpyxl", "text": "", "error": proc.stdout[:500]}
    return {key: str(data.get(key, "")) for key in ("status", "method", "text", "error")}


def image_metadata(path: Path) -> str:
    proc = run_command(["sips", "-g", "format", "-g", "pixelWidth", "-g", "pixelHeight", str(path)])
    if proc.returncode != 0:
        return f"Image file detected. OCR provider is not configured. sips error: {proc.stderr.strip()}"
    return "Image file detected. OCR provider is not configured.\n\n" + proc.stdout.strip()


def extract_file_text(path: Path) -> dict[str, str]:
    suffix = path.suffix.casefold()
    try:
        if suffix in TEXT_EXTENSIONS:
            return {"status": "complete", "method": "utf-8-text", "text": path.read_text(encoding="utf-8", errors="replace"), "error": ""}
        if suffix in HTML_EXTENSIONS:
            return {"status": "complete", "method": "html-parser", "text": strip_html(path.read_text(encoding="utf-8", errors="replace")), "error": ""}
        if suffix in RTF_EXTENSIONS:
            proc = run_command(["textutil", "-convert", "txt", "-stdout", str(path)])
            if proc.returncode != 0:
                return {"status": "failed", "method": "textutil", "text": "", "error": proc.stderr.strip() or proc.stdout.strip()}
            return {"status": "complete", "method": "textutil", "text": proc.stdout, "error": ""}
        if suffix in PDF_EXTENSIONS:
            return extract_pdf(path)
        if suffix in DOCX_EXTENSIONS:
            return extract_docx(path)
        if suffix in LEGACY_WORD_EXTENSIONS:
            return extract_word_with_textutil(path)
        if suffix in XLSX_EXTENSIONS:
            return extract_xlsx(path)
        if suffix in IMAGE_EXTENSIONS:
            ocr_result = ocr_images_with_vision([path], "vision-ocr-image")
            if ocr_result.get("status") == "complete":
                return ocr_result
            return {
                "status": "needs_ocr",
                "method": "image-metadata",
                "text": image_metadata(path),
                "error": ocr_result.get("error") or "OCR provider did not recognize text",
            }
    except OSError as exc:
        return {"status": "failed", "method": "filesystem", "text": "", "error": str(exc)}
    return {"status": "unsupported", "method": "extension", "text": "", "error": f"Unsupported file extension: {suffix or '<none>'}"}


def write_manifest(root: Path, manifest: dict[str, Any], touched: list[Path]) -> None:
    write_json(file_manifest_path(root, str(manifest["file_id"])), manifest, touched)


def concise_excerpt(text: str, limit: int = 1800) -> str:
    clean = re.sub(r"\n{3,}", "\n\n", text.strip())
    if len(clean) <= limit:
        return clean
    return clean[: limit - 3].rstrip() + "..."


def outline_from_text(title: str, file_id: str, extracted_rel: str, text: str) -> str:
    headings: list[str] = []
    for line in text.splitlines():
        clean = line.strip()
        if re.match(r"^#{1,6}\s+\S", clean):
            headings.append(clean)
        elif len(clean) <= 90 and clean.endswith(":") and len(clean.split()) <= 12:
            headings.append(clean)
    lines = [
        frontmatter(
            {
                "type": "personal_file_outline",
                "file_id": file_id,
                "updated": utc_now().isoformat().replace("+00:00", "Z"),
                "source": extracted_rel,
            }
        ).rstrip(),
        "",
        f"# Outline: {title}",
        "",
    ]
    if headings:
        lines.append("## Detected Structure")
        lines.append("")
        for heading in headings[:120]:
            lines.append(f"- {heading.lstrip('#').strip()}")
    else:
        lines.extend(
            [
                "## Detected Structure",
                "",
                "- No explicit headings detected. Use the extracted text for full inspection.",
            ]
        )
    lines.append("")
    return "\n".join(lines)


def summary_from_manifest(root: Path, manifest: dict[str, Any], text: str) -> str:
    file_id = str(manifest["file_id"])
    copied_rel = relative_or_empty(root, manifest.get("copied_path"))
    source_rel = relative_or_empty(root, manifest.get("source_path"))
    manifest_rel = relative(root, file_manifest_path(root, file_id))
    extracted_rel = relative(root, file_extracted_path(root, file_id)) if file_extracted_path(root, file_id).exists() else ""
    lines = [
        frontmatter(
            {
                "type": "personal_file_summary",
                "file_id": file_id,
                "title": manifest.get("title", file_id),
                "summary_status": manifest.get("summary_status", "pending_agent"),
                "extraction_status": manifest.get("extraction_status", "pending"),
                "privacy": "private",
                "updated": utc_now().isoformat().replace("+00:00", "Z"),
                "sources": [manifest_rel, copied_rel or source_rel],
            }
        ).rstrip(),
        "",
        f"# {manifest.get('title', file_id)}",
        "",
        "## Source",
        "",
        f"- Manifest: `{manifest_rel}`",
        f"- Original copy: `{copied_rel or 'external-only'}`",
        f"- Original source path: `{source_rel}`",
        f"- Extraction: `{manifest.get('extraction_status', 'pending')}` via `{manifest.get('extraction_method', '')}`",
        "",
        "## Summary Status",
        "",
        "This is a deterministic intake summary. A model-authored synthesis is still `pending_agent` unless explicitly reviewed in a Codex session.",
        "",
        "## Extracted Excerpt",
        "",
    ]
    if text.strip():
        lines.append(concise_excerpt(text))
    else:
        lines.append("No extracted text is available.")
    lines.append("")
    if extracted_rel:
        lines.extend(["## Full Extracted Text", "", f"- `{extracted_rel}`", ""])
    return "\n".join(lines)


def file_catalog_view(root: Path, manifests: list[dict[str, Any]]) -> str:
    lines = [
        frontmatter({"type": "view", "view": "file-catalog", "updated": utc_now().isoformat().replace("+00:00", "Z"), "file_count": len(manifests)}).rstrip(),
        "",
        "# File Catalog",
        "",
    ]
    by_kind: dict[str, list[dict[str, Any]]] = {}
    for manifest in manifests:
        by_kind.setdefault(str(manifest.get("kind", "other")), []).append(manifest)
    if not manifests:
        lines.append("No Personal OS files have been added yet.")
        lines.append("")
        return "\n".join(lines)
    for kind in sorted(by_kind):
        lines.extend([f"## {kind.title()}", ""])
        for manifest in by_kind[kind]:
            file_id = str(manifest.get("file_id", ""))
            title = str(manifest.get("title", file_id))
            summary = relative(root, file_summary_path(root, file_id)) if file_summary_path(root, file_id).exists() else ""
            manifest_rel = relative(root, file_manifest_path(root, file_id))
            bits = [
                f"`{manifest.get('extraction_status', 'pending')}` extraction",
                f"`{manifest.get('summary_status', 'pending')}` summary",
                f"added {manifest.get('added_at', '')}",
            ]
            target = summary or manifest_rel
            lines.append(f"- [{title}]({target}) - {'; '.join(bits)}")
        lines.append("")
    return "\n".join(lines)


def review_needed_rows(manifests: list[dict[str, Any]]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for manifest in manifests:
        reasons = []
        source_path = str(manifest.get("source_path", ""))
        generated_seed_artifact = "/personal-os-" in source_path and "-seed/" in source_path
        if manifest.get("extraction_status") in {"pending", "failed", "needs_ocr", "unsupported"}:
            reasons.append(str(manifest.get("extraction_status")))
        if manifest.get("summary_status") in {"pending", "failed"}:
            reasons.append(f"summary:{manifest.get('summary_status')}")
        if manifest.get("duplicate_of") and not generated_seed_artifact:
            reasons.append(f"duplicate_of:{manifest.get('duplicate_of')}")
        if manifest.get("source_changed"):
            reasons.append("source_changed")
        if reasons:
            rows.append({"manifest": manifest, "reasons": reasons})
    return rows


def review_needed_view(root: Path, manifests: list[dict[str, Any]]) -> str:
    rows = review_needed_rows(manifests)
    lines = [
        frontmatter({"type": "view", "view": "review-needed", "updated": utc_now().isoformat().replace("+00:00", "Z"), "item_count": len(rows)}).rstrip(),
        "",
        "# Review Needed",
        "",
    ]
    if not rows:
        lines.extend(["No file review items are currently open.", ""])
        return "\n".join(lines)
    for row in rows:
        manifest = row["manifest"]
        file_id = str(manifest.get("file_id", ""))
        title = str(manifest.get("title", file_id))
        manifest_rel = relative(root, file_manifest_path(root, file_id))
        lines.append(f"- [{title}]({manifest_rel}) - {', '.join(row['reasons'])}")
    lines.append("")
    return "\n".join(lines)


def manifest_domain(manifest: dict[str, Any]) -> str:
    paths = manifest.get("source_paths")
    source_path = ""
    if isinstance(paths, list) and paths:
        source_path = str(paths[0])
    if not source_path:
        source_path = str(manifest.get("source_path", ""))
    parts = Path(source_path).parts
    if "[Documents]" in parts:
        index = parts.index("[Documents]")
        if index + 1 < len(parts):
            return parts[index + 1]
    return "(unknown)"


def model_summary_progress_view(root: Path, manifests: list[dict[str, Any]]) -> str:
    by_status: dict[str, int] = {}
    by_domain: dict[str, dict[str, int]] = {}
    for manifest in manifests:
        status = str(manifest.get("summary_status", "pending"))
        by_status[status] = by_status.get(status, 0) + 1
        domain = manifest_domain(manifest)
        by_domain.setdefault(domain, {})
        by_domain[domain][status] = by_domain[domain].get(status, 0) + 1

    lines = [
        frontmatter(
            {
                "type": "view",
                "view": "model-summary-progress",
                "updated": utc_now().isoformat().replace("+00:00", "Z"),
                "file_count": len(manifests),
                "complete_count": by_status.get("complete", 0),
                "pending_agent_count": by_status.get("pending_agent", 0),
                "failed_count": by_status.get("failed", 0),
            }
        ).rstrip(),
        "",
        "# Model Summary Progress",
        "",
        "This view tracks files whose summaries have been model-authored and reviewed in a Codex session. Deterministic intake summaries remain `pending_agent`.",
        "",
        "## Totals",
        "",
    ]
    if not manifests:
        lines.append("- No files tracked.")
    else:
        for status in sorted(by_status):
            lines.append(f"- `{status}`: {by_status[status]}")
    lines.extend(["", "## By Domain", ""])
    for domain in sorted(by_domain):
        bits = [f"`{status}` {count}" for status, count in sorted(by_domain[domain].items())]
        lines.append(f"- {domain}: {'; '.join(bits)}")
    lines.extend(["", "## Recently Completed", ""])
    completed = [
        manifest
        for manifest in manifests
        if manifest.get("summary_status") == "complete" and manifest.get("model_summary_completed_at")
    ]
    completed.sort(key=lambda item: str(item.get("model_summary_completed_at", "")), reverse=True)
    if completed:
        for manifest in completed[:40]:
            file_id = str(manifest.get("file_id", ""))
            title = str(manifest.get("title", file_id))
            lines.append(
                f"- [{title}]({relative(root, file_summary_path(root, file_id))}) - completed {manifest.get('model_summary_completed_at', '')}"
            )
    else:
        lines.append("- None yet.")
    lines.extend(["", "## Remaining Queue", ""])
    remaining = [manifest for manifest in manifests if manifest.get("summary_status") != "complete"]
    if remaining:
        for manifest in remaining[:80]:
            file_id = str(manifest.get("file_id", ""))
            title = str(manifest.get("title", file_id))
            lines.append(
                f"- [{title}]({relative(root, file_manifest_path(root, file_id))}) - `{manifest.get('summary_status', 'pending')}` summary; `{manifest.get('extraction_status', 'pending')}` extraction"
            )
        if len(remaining) > 80:
            lines.append(f"- ... {len(remaining) - 80} more")
    else:
        lines.append("- No remaining files.")
    lines.append("")
    return "\n".join(lines)


def butlers_book_view(root: Path, manifests: list[dict[str, Any]]) -> str:
    recent = manifests[:20]
    things_state = load_json(things_state_path(root), [])
    open_tasks = [row for row in things_state if isinstance(row, dict) and not row.get("rolled_back")]
    lines = [
        frontmatter({"type": "view", "view": "butlers-book", "updated": utc_now().isoformat().replace("+00:00", "Z"), "privacy": "private"}).rstrip(),
        "",
        "# Personal OS Butler's Book",
        "",
        "This is the main agent entrypoint for private personal context. It links to source manifests and summaries rather than copying raw document text into the view.",
        "",
        "## Profile",
        "",
    ]
    for kind, filename in PROFILE_FILES.items():
        path = root / "profile" / filename
        if path.exists():
            lines.append(f"- [{kind.title()}](profile/{filename})")
    extra_profiles = sorted((root / "profile").glob("*.md"))
    for path in extra_profiles:
        if path.name not in PROFILE_FILES.values():
            title = path.stem.replace("-", " ").title()
            lines.append(f"- [{title}]({relative(root, path)})")
    lines.extend(["", "## Recent Notes", ""])
    recent_notes = [note for note in collect_notes(root) if note.get("kind") in {"journal", "episodes", "reflections", "profile"}][:12]
    if recent_notes:
        for note in recent_notes:
            lines.append(f"- [{note['title']}]({note['path']}) - `{note['kind']}`; updated {note['mtime']}")
    else:
        lines.append("- No notes captured yet.")
    lines.extend(["", "## Important Files", ""])
    if recent:
        for manifest in recent:
            file_id = str(manifest.get("file_id", ""))
            title = str(manifest.get("title", file_id))
            target = relative(root, file_summary_path(root, file_id)) if file_summary_path(root, file_id).exists() else relative(root, file_manifest_path(root, file_id))
            lines.append(f"- [{title}]({target}) - `{manifest.get('kind', 'other')}`; added {manifest.get('added_at', '')}")
    else:
        lines.append("- No files added yet.")
    lines.extend(["", "## Review Queues", ""])
    lines.append("- [Daily time plan](_views/time-daily.md)")
    lines.append("- [Weekly time plan](_views/time-weekly.md)")
    lines.append("- [Agent task candidates](_views/agent-task-candidates.md)")
    lines.append("- [File catalog](_views/file-catalog.md)")
    lines.append("- [Model summary progress](_views/model-summary-progress.md)")
    lines.append("- [Review needed](_views/review-needed.md)")
    lines.append("- [Wiki promotion candidates](_views/wiki-promotion-candidates.md)")
    lines.append("- [Calendar candidates](_views/calendar-candidates.md)")
    lines.extend(["", "## Agent-Created Things Follow-Ups", ""])
    if open_tasks:
        for row in open_tasks[-30:]:
            lines.append(f"- {row.get('title', row.get('id', 'Untitled task'))} (`{row.get('id', '')}`)")
    else:
        lines.append("- None recorded.")
    lines.append("")
    return "\n".join(lines)


def file_index_data(root: Path, manifests: list[dict[str, Any]]) -> dict[str, Any]:
    return {
        "type": "personal_os_file_index",
        "updated": utc_now().isoformat().replace("+00:00", "Z"),
        "file_count": len(manifests),
        "files": [
            {
                "file_id": manifest.get("file_id"),
                "title": manifest.get("title"),
                "kind": manifest.get("kind"),
                "sensitivity": manifest.get("sensitivity"),
                "manifest": relative(root, file_manifest_path(root, str(manifest.get("file_id")))),
                "summary": relative(root, file_summary_path(root, str(manifest.get("file_id"))))
                if file_summary_path(root, str(manifest.get("file_id"))).exists()
                else "",
                "extraction_status": manifest.get("extraction_status"),
                "summary_status": manifest.get("summary_status"),
                "added_at": manifest.get("added_at"),
            }
            for manifest in manifests
        ],
    }


def rebuild_file_views(root: Path, touched: list[Path]) -> None:
    manifests = load_file_manifests(root)
    write_json(root / "files" / "indexes" / "files-index.json", file_index_data(root, manifests), touched)
    write_text(root / "_views" / "file-catalog.md", file_catalog_view(root, manifests), touched)
    write_text(root / "_views" / "model-summary-progress.md", model_summary_progress_view(root, manifests), touched)
    write_text(root / "_views" / "review-needed.md", review_needed_view(root, manifests), touched)
    write_text(root / "_views" / "butlers-book.md", butlers_book_view(root, manifests), touched)
    if not (root / "_views" / "calendar-candidates.md").exists():
        write_text(root / "_views" / "calendar-candidates.md", calendar_candidates_stub(root), touched)
    if not (root / "_views" / "wiki-promotion-candidates.md").exists():
        write_text(root / "_views" / "wiki-promotion-candidates.md", wiki_promotion_stub(root), touched)


def start_run(args: argparse.Namespace, command: str, root: Path) -> dict[str, Any]:
    return {
        "run_id": run_id(),
        "command": command,
        "root": str(root),
        "started_at": utc_now().isoformat().replace("+00:00", "Z"),
        "finished_at": None,
        "git_commit_before": git_head(root),
        "git_commit_after": None,
        "files_touched": [],
        "things_created": [],
        "things_rolled_back": [],
        "automation": [],
        "rollback": [],
        "errors": [],
    }


def finish_run(
    args: argparse.Namespace,
    root: Path,
    record: dict[str, Any],
    touched: list[Path],
    commit_message: str,
) -> dict[str, Any]:
    record["finished_at"] = utc_now().isoformat().replace("+00:00", "Z")
    log_path = root / "_logs" / "runs" / f"{record['run_id']}.json"
    record["files_touched"] = sorted({relative(root, path) for path in [*touched, log_path]})
    log_path.parent.mkdir(parents=True, exist_ok=True)
    log_path.write_text(json.dumps(record, indent=2, ensure_ascii=False, sort_keys=True) + "\n", encoding="utf-8")
    content_commit = commit_all(root, commit_message, no_commit=args.no_commit)
    record["git_commit_after"] = content_commit
    log_path.write_text(json.dumps(record, indent=2, ensure_ascii=False, sort_keys=True) + "\n", encoding="utf-8")
    if not args.no_commit:
        commit_all(root, f"personal-os: log {record['run_id']}", no_commit=False)
    return record


def launch_agent_plist(root: Path, disabled: bool = True) -> dict[str, Any]:
    script = Path(__file__).resolve()
    return {
        "Label": LAUNCH_AGENT_LABEL,
        "ProgramArguments": [
            sys.executable,
            str(script),
            "--root",
            str(root),
            "garden",
            "daily",
            "--date",
            "today",
            "--write",
        ],
        "StartCalendarInterval": {"Hour": 22, "Minute": 30},
        "RunAtLoad": False,
        "Disabled": disabled,
        "StandardOutPath": str(root / "_logs" / "automation.out.log"),
        "StandardErrorPath": str(root / "_logs" / "automation.err.log"),
        "WorkingDirectory": str(root),
    }


def write_plist(path: Path, data: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("wb") as handle:
        plistlib.dump(data, handle, sort_keys=True)


def seed_rules() -> str:
    return """---
type: rules
version: 1
privacy: private
---

# Personal OS Rules

Personal OS is private by default. Raw journal entries and episodes stay here.
The external LLM wiki receives only explicit promotion candidates. Codex memory
under `~/.codex/memories` is operational memory and must not be used as the
personal journal.

## Reversibility

- Side effects require explicit commands.
- File writes are committed to this local git repo.
- Things writes are recorded in `_state/things-created.json`.
- Calendar writes are not allowed in v1.
- LaunchAgent automation is opt-in and reversible.

## Profile Notes

Profile notes may be updated automatically only from explicit profile candidate
lines in reflections. Every appended claim must cite a source file path.

## Files And Butler's Book

- Important personal files live under `files/`, not in `~/.codex/wiki`.
- Raw copied files live in `files/originals/` and are committed to this private
  local git repo by default.
- File manifests are the source of truth for document metadata, provenance,
  extraction status, and review state.
- Generated views under `_views/` are rebuildable scan surfaces.
- Missing OCR/model capability must be surfaced as `needs_ocr` or
  `pending_agent`, not hidden.

## Things Routing

Agent-created Personal OS follow-ups go to the Things project `🤖 Agent Tasks`
with the `Personal OS` marker. Writes must be recorded in
`_state/things-created.json` and rollback must affect only recorded tasks.

## Time And Next Actions

- Things is the master source for concrete next actions.
- Calendar is read-only context for time availability and deadline pressure.
- Personal OS may generate time reports under `_views/`, `_state/time/`, and
  `_reports/time/`; these are scan surfaces, not a second task manager.
- Agent-facing follow-ups from time reports must be explicit action lines and
  must route through the `🤖 Agent Tasks` project.
- External wiki/Raindrop experiment queues remain separate unless a reviewed
  item is promoted into Things.
"""


def seed_agents() -> str:
    return frontmatter(
        {
            "type": "agent_instructions",
            "privacy": "private",
            "version": 1,
        }
    ) + """
# Personal OS Agent Instructions

This folder contains private personal memory. Read it only for Personal OS,
journaling, reflection, or user-authorized introspection tasks.

Do not copy raw personal content into `~/.codex/wiki` or `~/.codex/memories`.
Use `~/.codex/skills/personal-os/scripts/personal.py` for writes.

Start with `_views/butlers-book.md` when you need private personal context.
Use `files/manifests/` for provenance and `files/summaries/` or
`files/extracted/` for progressive disclosure into specific documents.
Use `_views/time-daily.md` and `_views/time-weekly.md` for current time and
next-action scan surfaces when they exist.
"""


def profile_seed(kind: str) -> str:
    title = PROFILE_FILES[kind].replace(".md", "").replace("-", " ").title()
    return frontmatter(
        {
            "type": "profile_note",
            "status": "reviewed",
            "kind": kind,
            "privacy": "private",
            "created": today(),
            "updated": today(),
            "sources": [],
        }
    ) + f"\n# {title}\n\n## Entries\n\n"


def collect_notes(root: Path) -> list[dict[str, Any]]:
    notes: list[dict[str, Any]] = []
    for base in ("journal", "episodes", "reflections", "profile", "files"):
        for path in sorted((root / base).glob("**/*.md")):
            try:
                text = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                continue
            stat = path.stat()
            title_match = re.search(r"(?m)^#\s+(.+)$", text)
            notes.append(
                {
                    "path": relative(root, path),
                    "title": title_match.group(1).strip() if title_match else path.stem,
                    "mtime": dt.datetime.fromtimestamp(stat.st_mtime, UTC).isoformat().replace("+00:00", "Z"),
                    "kind": base,
                }
            )
    notes.sort(key=lambda row: row["mtime"], reverse=True)
    return notes


def render_index(root: Path) -> str:
    notes = collect_notes(root)
    lines = [
        frontmatter({"type": "index", "updated": utc_now().isoformat().replace("+00:00", "Z"), "page_count": len(notes)}).rstrip(),
        "",
        "# Personal OS Index",
        "",
        "## Recent Notes",
        "",
    ]
    for note in notes[:120]:
        lines.append(f"- [{note['title']}]({note['path']}) - `{note['kind']}`; updated {note['mtime']}")
    lines.append("")
    return "\n".join(lines)


def cmd_bootstrap(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    root.mkdir(parents=True, exist_ok=True)
    for name in REQUIRED_DIRS:
        (root / name).mkdir(parents=True, exist_ok=True)
    ensure_private_git(root)
    record = start_run(args, "bootstrap", root)
    touched: list[Path] = []

    seeds = {
        root / "rules.md": seed_rules(),
        root / "AGENTS.md": seed_agents(),
        root / "index.md": render_index(root),
        root / ".gitignore": "_logs/*.log\n.DS_Store\n",
    }
    for path, text in seeds.items():
        if not path.exists():
            write_text(path, text, touched)

    append_section_if_missing(
        root / "rules.md",
        "## Files And Butler's Book",
        """
## Files And Butler's Book

- Important personal files live under `files/`, not in `~/.codex/wiki`.
- Raw copied files live in `files/originals/` and are committed to this private
  local git repo by default.
- File manifests are the source of truth for document metadata, provenance,
  extraction status, and review state.
- Generated views under `_views/` are rebuildable scan surfaces.
- Missing OCR/model capability must be surfaced as `needs_ocr` or
  `pending_agent`, not hidden.

## Things Routing

Agent-created Personal OS follow-ups go to the Things project `🤖 Agent Tasks`
with the `Personal OS` marker. Writes must be recorded in
`_state/things-created.json` and rollback must affect only recorded tasks.
""",
        touched,
    )
    append_section_if_missing(
        root / "rules.md",
        "## Time And Next Actions",
        """
## Time And Next Actions

- Things is the master source for concrete next actions.
- Calendar is read-only context for time availability and deadline pressure.
- Personal OS may generate time reports under `_views/`, `_state/time/`, and
  `_reports/time/`; these are scan surfaces, not a second task manager.
- Agent-facing follow-ups from time reports must be explicit action lines and
  must route through the `🤖 Agent Tasks` project.
- External wiki/Raindrop experiment queues remain separate unless a reviewed
  item is promoted into Things.
""",
        touched,
    )
    append_section_if_missing(
        root / "AGENTS.md",
        "_views/butlers-book.md",
        """
Start with `_views/butlers-book.md` when you need private personal context.
Use `files/manifests/` for provenance and `files/summaries/` or
`files/extracted/` for progressive disclosure into specific documents.
""",
        touched,
    )

    for kind, filename in PROFILE_FILES.items():
        path = root / "profile" / filename
        if not path.exists():
            write_text(path, profile_seed(kind), touched)

    routing_path = things_routing_path(root)
    if not routing_path.exists():
        write_json(routing_path, default_things_routing(), touched)

    template_path = root / "_state" / "launchagents" / f"{LAUNCH_AGENT_LABEL}.plist"
    write_plist(template_path, launch_agent_plist(root, disabled=True))
    touched.append(template_path)
    rebuild_file_views(root, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append("Remove ~/.codex/personal-os only if you want to delete the private Personal OS root.")
    finish_run(args, root, record, touched, "personal-os: bootstrap")
    print_json({"root": str(root), "run_id": record["run_id"], "status": "bootstrapped"})


def journal_path(root: Path, day: dt.date) -> Path:
    return root / "journal" / str(day.year) / f"{day.year}-{day.month:02d}" / f"{day.isoformat()}.md"


def episode_path(root: Path, day: dt.date, title: str) -> Path:
    return root / "episodes" / str(day.year) / f"{day.isoformat()}-{slugify(title)}.md"


def cmd_capture_journal(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    text = read_stdin(args)
    day = parse_day(args.date)
    path = journal_path(root, day)
    record = start_run(args, "capture journal", root)
    touched: list[Path] = []
    if not path.exists():
        body = frontmatter(
            {
                "type": "journal_entry",
                "date": day.isoformat(),
                "privacy": "private",
                "created": utc_now().isoformat().replace("+00:00", "Z"),
                "updated": utc_now().isoformat().replace("+00:00", "Z"),
                "tags": [],
            }
        ) + f"\n# Journal {day.isoformat()}\n\n"
        write_text(path, body, touched)
    append_text(path, f"## {args.title}\n\n{text}\n", touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for the commit, or edit {relative(root, path)} to remove this journal section.")
    finish_run(args, root, record, touched, f"personal-os: capture journal {day.isoformat()}")
    print_json({"path": str(path), "run_id": record["run_id"]})


def cmd_capture_episode(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    text = read_stdin(args)
    day = parse_day(args.date)
    path = episode_path(root, day, args.title)
    record = start_run(args, "capture episode", root)
    touched: list[Path] = []
    body = frontmatter(
        {
            "type": "episode",
            "source": args.source,
            "date": day.isoformat(),
            "privacy": "private",
            "created": utc_now().isoformat().replace("+00:00", "Z"),
            "updated": utc_now().isoformat().replace("+00:00", "Z"),
        }
    ) + f"\n# {args.title}\n\n{text}\n"
    write_text(path, body, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for the commit, or remove {relative(root, path)}.")
    finish_run(args, root, record, touched, f"personal-os: capture episode {day.isoformat()}")
    print_json({"path": str(path), "run_id": record["run_id"]})


def read_day_inputs(root: Path, day: dt.date, include_chronicle: bool = True) -> list[dict[str, str]]:
    inputs: list[dict[str, str]] = []
    jpath = journal_path(root, day)
    if jpath.exists():
        inputs.append({"kind": "journal", "path": str(jpath), "text": jpath.read_text(encoding="utf-8")})
    for path in sorted((root / "episodes").glob(f"**/{day.isoformat()}-*.md")):
        inputs.append({"kind": "episode", "path": str(path), "text": path.read_text(encoding="utf-8")})
    if include_chronicle and CHRONICLE_RESOURCES.exists():
        for path in sorted(CHRONICLE_RESOURCES.glob(f"{day.isoformat()}T*.md")):
            try:
                text = path.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                continue
            inputs.append({"kind": "chronicle", "path": str(path), "text": text})
    return inputs


def excerpt(text: str, limit: int = 900) -> str:
    clean = re.sub(r"\s+", " ", text).strip()
    if len(clean) <= limit:
        return clean
    return clean[: limit - 3].rstrip() + "..."


def explicit_actions_from_text(text: str) -> list[str]:
    actions: list[str] = []
    seen: set[str] = set()
    for line in text.splitlines():
        match = ACTION_RE.match(line) or CHECKBOX_RE.match(line)
        if match:
            title = match.group(match.lastindex or 1).strip()
            key = re.sub(r"\s+", " ", title.casefold()).strip()
            if len(title) >= 3 and key not in seen:
                seen.add(key)
                actions.append(title)
    return actions


def profile_candidates_from_text(text: str) -> list[dict[str, str]]:
    candidates: list[dict[str, str]] = []
    for line in text.splitlines():
        match = PROFILE_RE.match(line)
        if match:
            candidates.append({"kind": match.group(1).casefold(), "text": match.group(2).strip()})
    return candidates


def render_daily_reflection(root: Path, day: dt.date, inputs: list[dict[str, str]]) -> str:
    source_paths = [relative(root, Path(item["path"])) if str(item["path"]).startswith(str(root)) else item["path"] for item in inputs]
    lines = [
        frontmatter(
            {
                "type": "reflection",
                "period_start": day.isoformat(),
                "period_end": day.isoformat(),
                "status": "draft",
                "privacy": "private",
                "created": utc_now().isoformat().replace("+00:00", "Z"),
                "updated": utc_now().isoformat().replace("+00:00", "Z"),
                "inputs": source_paths,
            }
        ).rstrip(),
        "",
        f"# Daily Reflection {day.isoformat()}",
        "",
        "## Inputs",
        "",
    ]
    if not inputs:
        lines.append("- No journal, episode, or Chronicle inputs found.")
    for item in inputs:
        path_text = relative(root, Path(item["path"])) if str(item["path"]).startswith(str(root)) else item["path"]
        lines.append(f"- `{item['kind']}`: {path_text}")
    lines.extend(["", "## Evidence", ""])
    for item in inputs:
        path_text = relative(root, Path(item["path"])) if str(item["path"]).startswith(str(root)) else item["path"]
        lines.append(f"### {item['kind']}: {path_text}")
        lines.append("")
        lines.append(excerpt(item["text"]))
        lines.append("")
    actions = []
    for item in inputs:
        actions.extend(explicit_actions_from_text(item["text"]))
    lines.extend(["## Explicit Action Candidates", ""])
    if actions:
        for action in actions:
            lines.append(f"- [ ] {action}")
    else:
        lines.append("- None found.")
    lines.extend(
        [
            "",
            "## Profile Update Candidates",
            "",
            "Add explicit lines here when appropriate:",
            "",
            "- preference: ...",
            "- pattern: ...",
            "- routine: ...",
            "- value: ...",
            "",
            "## Notes",
            "",
            "Draft generated from available inputs. Edit before treating it as reviewed.",
            "",
        ]
    )
    return "\n".join(lines)


def reflection_path(root: Path, day: dt.date) -> Path:
    return root / "reflections" / "daily" / f"{day.isoformat()}.md"


def cmd_reflect_daily(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    day = parse_day(args.date)
    inputs = read_day_inputs(root, day, include_chronicle=not args.no_chronicle)
    rendered = render_daily_reflection(root, day, inputs)
    if not args.write:
        print(rendered)
        return
    path = reflection_path(root, day)
    record = start_run(args, "reflect daily", root)
    touched: list[Path] = []
    write_text(path, rendered, touched)
    write_text(root / "_views" / "calendar-candidates.md", calendar_candidates_stub(root), touched)
    write_text(root / "_views" / "wiki-promotion-candidates.md", wiki_promotion_stub(root), touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for the commit, or remove {relative(root, path)}.")
    finish_run(args, root, record, touched, f"personal-os: daily reflection {day.isoformat()}")
    print_json({"path": str(path), "run_id": record["run_id"], "input_count": len(inputs)})


def calendar_candidates_stub(root: Path) -> str:
    return frontmatter({"type": "view", "view": "calendar-candidates", "updated": utc_now().isoformat().replace("+00:00", "Z")}) + "\n# Calendar Candidates\n\nCalendar writes are disabled in v1. Add reviewed candidates here manually.\n"


def wiki_promotion_stub(root: Path) -> str:
    return frontmatter({"type": "view", "view": "wiki-promotion-candidates", "updated": utc_now().isoformat().replace("+00:00", "Z")}) + "\n# Wiki Promotion Candidates\n\nNo raw personal content should be promoted. Add reviewed, non-sensitive synthesis candidates here.\n"


def cmd_profile_update(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    source = ensure_under_root(root, Path(args.from_reflection))
    text = source.read_text(encoding="utf-8")
    candidates = profile_candidates_from_text(text)
    if not args.apply:
        print_json({"source": relative(root, source), "candidates": candidates})
        return
    record = start_run(args, "profile update", root)
    touched: list[Path] = []
    by_file: dict[Path, list[str]] = {}
    for candidate in candidates:
        filename = PROFILE_FILES.get(candidate["kind"])
        if not filename:
            continue
        path = root / "profile" / filename
        line = f"- {candidate['text']} (source: `{relative(root, source)}`; added {today()})\n"
        by_file.setdefault(path, []).append(line)
    for path, lines in by_file.items():
        append_text(path, "".join(lines), touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append("Use git revert for the profile update commit.")
    finish_run(args, root, record, touched, "personal-os: profile update")
    print_json({"run_id": record["run_id"], "updated_files": [relative(root, path) for path in by_file]})


def action_plan_from_reflection(root: Path, source: Path) -> dict[str, Any]:
    text = source.read_text(encoding="utf-8")
    tasks = []
    for title in explicit_actions_from_text(text):
        tasks.append(
            {
                "title": title,
                "notes": f"Created from Personal OS reflection: {relative(root, source)}",
                "source": relative(root, source),
                "tag": "Personal OS",
            }
        )
    return {
        "type": "personal_os_action_plan",
        "created": utc_now().isoformat().replace("+00:00", "Z"),
        "source_reflection": relative(root, source),
        "tasks": tasks,
    }


def cmd_actions_plan(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    source = ensure_under_root(root, Path(args.from_reflection))
    plan = action_plan_from_reflection(root, source)
    if not args.write:
        print_json(plan)
        return
    record = start_run(args, "actions plan", root)
    touched: list[Path] = []
    plan_path = root / "_views" / f"action-plan-{record['run_id']}.json"
    write_json(plan_path, plan, touched)
    record["rollback"].append(f"Remove {relative(root, plan_path)} or use git revert.")
    finish_run(args, root, record, touched, "personal-os: action plan")
    print_json({"run_id": record["run_id"], "plan": str(plan_path), "task_count": len(plan["tasks"])})


def load_action_plan(root: Path, path: Path) -> dict[str, Any]:
    resolved = ensure_under_root(root, path)
    data = json.loads(resolved.read_text(encoding="utf-8"))
    if data.get("type") != "personal_os_action_plan":
        raise SystemExit(f"Not a Personal OS action plan: {path}")
    return data


def call_things(args: list[str]) -> Any:
    if not THINGS_SCRIPT.exists():
        raise SystemExit(f"Things helper not found: {THINGS_SCRIPT}")
    proc = run_command([str(THINGS_SCRIPT), *args], timeout=THINGS_TIMEOUT_SECONDS)
    if proc.returncode != 0:
        raise SystemExit(proc.stderr.strip() or proc.stdout.strip() or "Things helper failed")
    return json.loads(proc.stdout)


def verify_things_item(item_id: str) -> list[dict[str, Any]]:
    try:
        snapshot = call_things(["snapshot-db"])
    except SystemExit:
        snapshot = call_things(["snapshot"])
    return [todo for todo in snapshot.get("todos", []) if str(todo.get("id")) == str(item_id)]


def add_file_to_root(
    root: Path,
    source: Path,
    *,
    kind: str,
    sensitivity: str | None,
    title: str | None,
    notes: str,
    copy_file: bool,
    touched: list[Path],
) -> dict[str, Any]:
    source = source.expanduser().resolve()
    if not source.exists():
        raise SystemExit(f"File does not exist: {source}")
    if not source.is_file():
        raise SystemExit(f"Expected a file: {source}")
    if kind not in FILE_KINDS:
        raise SystemExit(f"Unsupported file kind {kind!r}. Use one of: {', '.join(FILE_KINDS)}")
    if sensitivity is not None and sensitivity not in SENSITIVITY_LABELS:
        raise SystemExit(f"Unsupported sensitivity {sensitivity!r}. Use one of: {', '.join(SENSITIVITY_LABELS)}")
    digest = sha256_file(source)
    title_text = file_title_from_path(source, title)
    file_id = file_id_for(source, title_text, digest)
    existing_path = file_manifest_path(root, file_id)
    existing = load_json(existing_path, {}) if existing_path.exists() else {}
    source_paths = []
    for value in existing.get("source_paths", []):
        if value and value not in source_paths:
            source_paths.append(str(value))
    previous_source = existing.get("source_path")
    if previous_source and previous_source not in source_paths:
        source_paths.append(str(previous_source))
    if str(source) not in source_paths:
        source_paths.append(str(source))
    copied_path = str(existing.get("copied_path") or "")
    if copy_file:
        dest_dir = root / "files" / "originals" / file_id
        dest_dir.mkdir(parents=True, exist_ok=True)
        dest = dest_dir / source.name
        if not dest.exists() or sha256_file(dest) != digest:
            shutil.copy2(source, dest)
            touched.append(dest)
        copied_path = str(dest)
    duplicate_of = existing.get("duplicate_of") or duplicate_for_sha(root, digest, file_id)
    now = utc_now().isoformat().replace("+00:00", "Z")
    manifest = {
        "file_id": file_id,
        "title": title_text,
        "kind": kind,
        "sensitivity": sensitivity or sensitivity_for_kind(kind),
        "source_path": str(source),
        "source_paths": source_paths,
        "copied_path": copied_path,
        "original_filename": source.name,
        "sha256": digest,
        "size_bytes": source.stat().st_size,
        "mime_type": detect_mime(source),
        "added_at": existing.get("added_at") or now,
        "updated_at": now,
        "source_modified_at": source_modified_at(source),
        "extraction_status": existing.get("extraction_status") or "pending",
        "summary_status": existing.get("summary_status") or "pending",
        "notes": notes,
        "tags": existing.get("tags") or [],
        "related_people": existing.get("related_people") or [],
        "related_places": existing.get("related_places") or [],
        "review_after": existing.get("review_after"),
        "supersedes": existing.get("supersedes"),
        "duplicate_of": duplicate_of,
        "source_exists": True,
        "source_changed": False,
    }
    for optional in ("extraction_method", "extracted_path", "extracted_at", "extraction_error", "summary_path", "outline_path", "summarized_at"):
        if optional in existing:
            manifest[optional] = existing[optional]
    write_manifest(root, manifest, touched)
    return manifest


def should_skip_bulk_path(path: Path) -> bool:
    return any(fnmatch.fnmatch(path.name, pattern) for pattern in BULK_EXCLUDE_PATTERNS)


def collect_folder_candidates(folder: Path, recursive: bool) -> dict[str, Any]:
    folder = folder.expanduser().resolve()
    if not folder.is_dir():
        raise SystemExit(f"Expected a folder: {folder}")
    candidates: list[dict[str, Any]] = []
    skipped: list[dict[str, str]] = []
    iterator: list[Path] = []
    if recursive:
        for current, dirs, files in os.walk(folder):
            current_path = Path(current)
            kept_dirs = []
            for dirname in dirs:
                dpath = current_path / dirname
                if should_skip_bulk_path(dpath):
                    skipped.append({"path": str(dpath), "reason": "excluded-directory"})
                else:
                    kept_dirs.append(dirname)
            dirs[:] = kept_dirs
            for filename in files:
                iterator.append(current_path / filename)
    else:
        iterator = [path for path in folder.iterdir() if path.is_file()]
        for path in folder.iterdir():
            if path.is_dir() and should_skip_bulk_path(path):
                skipped.append({"path": str(path), "reason": "excluded-directory"})
    for path in sorted(iterator):
        if should_skip_bulk_path(path):
            skipped.append({"path": str(path), "reason": "excluded-file"})
            continue
        try:
            size = path.stat().st_size
        except OSError as exc:
            skipped.append({"path": str(path), "reason": str(exc)})
            continue
        candidates.append(
            {
                "path": str(path),
                "size_bytes": size,
                "large": size >= LARGE_FILE_THRESHOLD_BYTES,
                "mime_type": detect_mime(path),
            }
        )
    return {"folder": str(folder), "candidates": candidates, "skipped": skipped}


def cmd_file_add(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    record = start_run(args, "file add", root)
    touched: list[Path] = []
    manifest = add_file_to_root(
        root,
        Path(args.path),
        kind=args.kind,
        sensitivity=args.sensitivity,
        title=args.title,
        notes=args.notes or "",
        copy_file=args.copy_mode == "copy",
        touched=touched,
    )
    rebuild_file_views(root, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for this commit, or remove files/manifests/{manifest['file_id']}.json and files/originals/{manifest['file_id']}/.")
    finish_run(args, root, record, touched, f"personal-os: add file {manifest['file_id']}")
    print_json({"run_id": record["run_id"], "manifest": manifest})


def cmd_file_add_folder(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    report = collect_folder_candidates(Path(args.path), recursive=args.recursive)
    if not args.write:
        print_json({**report, "dry_run": True, "write_required": "Re-run with --write to copy candidate files."})
        return
    record = start_run(args, "file add-folder", root)
    touched: list[Path] = []
    manifests = []
    for candidate in report["candidates"]:
        manifests.append(
            add_file_to_root(
                root,
                Path(candidate["path"]),
                kind=args.kind,
                sensitivity=args.sensitivity,
                title=None,
                notes=args.notes or "",
                copy_file=True,
                touched=touched,
            )
        )
    rebuild_file_views(root, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append("Use git revert for this commit to remove copied files and manifests from Personal OS.")
    finish_run(args, root, record, touched, f"personal-os: add folder {Path(args.path).name}")
    print_json({"run_id": record["run_id"], "added_count": len(manifests), "skipped": report["skipped"], "large_files": [row for row in report["candidates"] if row["large"]]})


def extract_manifest(root: Path, manifest: dict[str, Any], touched: list[Path]) -> dict[str, Any]:
    file_id = str(manifest["file_id"])
    source = manifest_source_path(manifest)
    if source is None or not source.exists():
        manifest["extraction_status"] = "failed"
        manifest["extraction_error"] = "No available copied or source file."
        manifest["extracted_at"] = utc_now().isoformat().replace("+00:00", "Z")
        write_manifest(root, manifest, touched)
        return manifest
    result = extract_file_text(source)
    status = result.get("status", "failed")
    if status not in EXTRACTION_STATUSES:
        status = "failed"
    extraction_path = file_extracted_path(root, file_id)
    extracted_text = result.get("text", "")
    if status != "complete" and result.get("error"):
        extracted_text = (extracted_text + "\n\n" if extracted_text else "") + f"Extraction status: {status}\nError: {result['error']}\n"
    write_text(extraction_path, extracted_text, touched)
    manifest["extraction_status"] = status
    manifest["extraction_method"] = result.get("method", "")
    manifest["extracted_path"] = str(extraction_path)
    manifest["extracted_at"] = utc_now().isoformat().replace("+00:00", "Z")
    manifest["extraction_error"] = result.get("error", "")
    write_manifest(root, manifest, touched)
    return manifest


def cmd_file_extract(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    manifest = find_manifest(root, args.file_id)
    if not args.write:
        print_json({"file_id": args.file_id, "would_extract_from": str(manifest_source_path(manifest) or ""), "write_required": True})
        return
    record = start_run(args, "file extract", root)
    touched: list[Path] = []
    manifest = extract_manifest(root, manifest, touched)
    rebuild_file_views(root, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for the extraction commit for {args.file_id}.")
    finish_run(args, root, record, touched, f"personal-os: extract file {args.file_id}")
    print_json({"run_id": record["run_id"], "file_id": args.file_id, "extraction_status": manifest.get("extraction_status")})


def summarize_manifest(root: Path, manifest: dict[str, Any], touched: list[Path]) -> dict[str, Any]:
    file_id = str(manifest["file_id"])
    if manifest.get("extraction_status") == "pending":
        manifest = extract_manifest(root, manifest, touched)
    extraction_path = file_extracted_path(root, file_id)
    text = extraction_path.read_text(encoding="utf-8", errors="replace") if extraction_path.exists() else ""
    if manifest.get("extraction_status") in {"failed", "unsupported"}:
        manifest["summary_status"] = "failed"
    else:
        manifest["summary_status"] = "pending_agent"
    summary = summary_from_manifest(root, manifest, text)
    outline = outline_from_text(str(manifest.get("title", file_id)), file_id, relative(root, extraction_path), text)
    write_text(file_summary_path(root, file_id), summary, touched)
    write_text(file_outline_path(root, file_id), outline, touched)
    manifest["summary_path"] = str(file_summary_path(root, file_id))
    manifest["outline_path"] = str(file_outline_path(root, file_id))
    manifest["summarized_at"] = utc_now().isoformat().replace("+00:00", "Z")
    write_manifest(root, manifest, touched)
    return manifest


def cmd_file_summarize(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    manifest = find_manifest(root, args.file_id)
    if not args.write:
        print_json({"file_id": args.file_id, "would_summarize": True, "current_summary_status": manifest.get("summary_status"), "write_required": True})
        return
    record = start_run(args, "file summarize", root)
    touched: list[Path] = []
    manifest = summarize_manifest(root, manifest, touched)
    rebuild_file_views(root, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for the summary commit for {args.file_id}.")
    finish_run(args, root, record, touched, f"personal-os: summarize file {args.file_id}")
    print_json({"run_id": record["run_id"], "file_id": args.file_id, "summary_status": manifest.get("summary_status"), "summary": str(file_summary_path(root, args.file_id))})


def cmd_file_inspect(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    manifest = find_manifest(root, args.file_id)
    file_id = str(manifest["file_id"])
    print_json(
        {
            "manifest": manifest,
            "paths": {
                "manifest": str(file_manifest_path(root, file_id)),
                "extracted": str(file_extracted_path(root, file_id)) if file_extracted_path(root, file_id).exists() else "",
                "summary": str(file_summary_path(root, file_id)) if file_summary_path(root, file_id).exists() else "",
                "outline": str(file_outline_path(root, file_id)) if file_outline_path(root, file_id).exists() else "",
            },
        }
    )


def snippet_for_query(text: str, query: str, limit: int = 280) -> str:
    folded = text.casefold()
    index = folded.find(query.casefold())
    if index < 0:
        return ""
    start = max(0, index - limit // 2)
    end = min(len(text), index + limit // 2)
    return re.sub(r"\s+", " ", text[start:end]).strip()


def cmd_file_search(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    query = args.query.casefold()
    hits = []
    for manifest in load_file_manifests(root):
        file_id = str(manifest.get("file_id", ""))
        haystacks = [
            ("manifest", json.dumps(manifest, ensure_ascii=False)),
        ]
        for label, path in (("summary", file_summary_path(root, file_id)), ("extracted", file_extracted_path(root, file_id))):
            if path.exists():
                haystacks.append((label, path.read_text(encoding="utf-8", errors="replace")))
        for label, text in haystacks:
            if query in text.casefold():
                hits.append(
                    {
                        "file_id": file_id,
                        "title": manifest.get("title"),
                        "kind": manifest.get("kind"),
                        "match": label,
                        "snippet": snippet_for_query(text, args.query),
                        "manifest": str(file_manifest_path(root, file_id)),
                    }
                )
                break
    print_json({"query": args.query, "count": len(hits), "hits": hits[: args.limit]})


def recheck_manifests(root: Path, touched: list[Path], write: bool) -> list[dict[str, Any]]:
    updates = []
    for manifest in load_file_manifests(root):
        changed = False
        source_value = manifest.get("source_path")
        copied_value = manifest.get("copied_path")
        if source_value:
            source = Path(str(source_value)).expanduser()
            source_exists = source.exists()
            current_sha = sha256_file(source) if source_exists and source.is_file() else ""
            source_changed = bool(current_sha and current_sha != manifest.get("sha256"))
            if manifest.get("source_exists") != source_exists or manifest.get("source_current_sha") != current_sha or manifest.get("source_changed") != source_changed:
                manifest["source_exists"] = source_exists
                manifest["source_current_sha"] = current_sha
                manifest["source_changed"] = source_changed
                manifest["source_checked_at"] = utc_now().isoformat().replace("+00:00", "Z")
                changed = True
        if copied_value:
            copied = Path(str(copied_value)).expanduser()
            copied_exists = copied.exists()
            copied_sha = sha256_file(copied) if copied_exists and copied.is_file() else ""
            copied_changed = bool(copied_sha and copied_sha != manifest.get("sha256"))
            if manifest.get("copied_exists") != copied_exists or manifest.get("copied_current_sha") != copied_sha or manifest.get("copied_changed") != copied_changed:
                manifest["copied_exists"] = copied_exists
                manifest["copied_current_sha"] = copied_sha
                manifest["copied_changed"] = copied_changed
                changed = True
        if changed:
            updates.append(manifest)
            if write:
                write_manifest(root, manifest, touched)
    return updates


def cmd_file_recheck(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    if not args.write:
        touched: list[Path] = []
        print_json({"updates": recheck_manifests(root, touched, write=False), "write_required": True})
        return
    record = start_run(args, "file recheck", root)
    touched = []
    updates = recheck_manifests(root, touched, write=True)
    rebuild_file_views(root, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append("Use git revert for the recheck commit.")
    finish_run(args, root, record, touched, "personal-os: recheck files")
    print_json({"run_id": record["run_id"], "updated_count": len(updates)})


def verify_file_layer(root: Path) -> list[str]:
    errors: list[str] = []
    for name in ("files/originals", "files/manifests", "files/extracted", "files/summaries", "files/outlines", "files/indexes"):
        if not (root / name).is_dir():
            errors.append(f"Missing file-layer directory: {name}")
    routing_file = things_routing_path(root)
    routing = load_json(routing_file, None)
    if routing is None:
        errors.append(f"Missing Things routing file: {relative(root, routing_file)}")
    else:
        for field in ("destination", "project_name", "project_id", "tag"):
            if field not in routing:
                errors.append(f"Things routing missing {field}")
    for path in sorted((root / "files" / "manifests").glob("*.json")):
        try:
            manifest = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            errors.append(f"Invalid file manifest JSON: {relative(root, path)}")
            continue
        for field in ("file_id", "title", "kind", "sensitivity", "source_path", "copied_path", "original_filename", "sha256", "size_bytes", "mime_type", "added_at", "source_modified_at", "extraction_status", "summary_status", "notes", "tags", "related_people", "related_places", "review_after", "supersedes", "duplicate_of"):
            if field not in manifest:
                errors.append(f"Manifest missing {field}: {relative(root, path)}")
        if manifest.get("kind") not in FILE_KINDS:
            errors.append(f"Manifest has unsupported kind: {relative(root, path)}")
        if manifest.get("extraction_status") not in EXTRACTION_STATUSES:
            errors.append(f"Manifest has invalid extraction_status: {relative(root, path)}")
        if manifest.get("summary_status") not in SUMMARY_STATUSES:
            errors.append(f"Manifest has invalid summary_status: {relative(root, path)}")
        if "source_paths" in manifest and not isinstance(manifest.get("source_paths"), list):
            errors.append(f"Manifest source_paths must be a list: {relative(root, path)}")
        copied = str(manifest.get("copied_path") or "")
        if copied and not Path(copied).expanduser().exists():
            errors.append(f"Copied file missing for manifest: {relative(root, path)}")
    return errors


def cmd_file_verify(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    errors = verify_file_layer(root)
    print_json({"root": str(root), "ok": not errors, "errors": errors, "file_count": len(load_file_manifests(root))})
    if errors:
        raise SystemExit(1)


def things_effective_status(item: dict[str, Any]) -> str:
    if item.get("cancellationDate"):
        return "canceled"
    if item.get("completionDate"):
        return "completed"
    return str(item.get("status") or "unknown")


def things_state_path(root: Path) -> Path:
    return root / "_state" / "things-created.json"


def verify_agent_tasks_project(root: Path) -> dict[str, Any]:
    routing = load_things_routing(root)
    data = call_things(["lists"])
    for project in data.get("projects", []):
        if project.get("id") == routing["project_id"] and project.get("name") == routing["project_name"]:
            return project
    by_name = [project for project in data.get("projects", []) if project.get("name") == routing["project_name"]]
    if by_name:
        raise SystemExit(
            f"Things project {routing['project_name']!r} exists but id differs. "
            f"Expected {routing['project_id']}, found {[project.get('id') for project in by_name]}."
        )
    raise SystemExit(f"Things project not found: {routing['project_name']}")


def active_things_source_keys(state: list[Any]) -> set[str]:
    keys = set()
    for row in state:
        if isinstance(row, dict) and not row.get("rolled_back") and row.get("source_key"):
            keys.add(str(row["source_key"]))
    return keys


def create_personal_os_things_task(
    root: Path,
    *,
    title: str,
    notes: str,
    source: str,
    source_key: str,
    run_id_value: str,
    state: list[Any],
) -> dict[str, Any] | None:
    if source_key in active_things_source_keys(state):
        return None
    routing = load_things_routing(root)
    verify_agent_tasks_project(root)
    result = call_things(
        [
            "add",
            title,
            "--notes",
            f"{notes}\n\nPersonal OS source: {source}\nPersonal OS run: {run_id_value}",
            "--tag",
            routing["tag"],
            "--list",
            routing["project_name"],
        ]
    )
    item_id = result.get("id") or result.get("uuid") or result.get("identifier")
    created_row = {
        "id": item_id,
        "title": result.get("name", title),
        "source": source,
        "source_key": source_key,
        "project_name": routing["project_name"],
        "project_id": routing["project_id"],
        "tag": routing["tag"],
        "run_id": run_id_value,
        "created_at": utc_now().isoformat().replace("+00:00", "Z"),
        "rolled_back": False,
    }
    if item_id:
        created_row["creation_verification"] = verify_things_item(str(item_id))
        if created_row["creation_verification"]:
            created_row["creation_effective_status"] = things_effective_status(created_row["creation_verification"][0])
    state.append(created_row)
    return created_row


def cmd_actions_apply(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    plan_path = ensure_under_root(root, Path(args.plan))
    plan = load_action_plan(root, plan_path)
    record = start_run(args, "actions apply", root)
    touched: list[Path] = []
    state = load_json(things_state_path(root), [])
    created = []
    verify_agent_tasks_project(root)
    for task in plan.get("tasks", []):
        title = str(task.get("title", "")).strip()
        if not title:
            continue
        notes = f"{task.get('notes', '')}\n\nPersonal OS run: {record['run_id']}"
        source = str(task.get("source", plan.get("source_reflection", "")))
        source_key = str(task.get("source_key") or f"action:{source}:{slugify(title)}")
        created_row = create_personal_os_things_task(
            root,
            title=title,
            notes=notes,
            source=source,
            source_key=source_key,
            run_id_value=record["run_id"],
            state=state,
        )
        if created_row:
            created_row["plan"] = relative(root, plan_path)
            created.append(created_row)
    write_json(things_state_path(root), state, touched)
    record["things_created"] = created
    record["rollback"].append(f"Run `personal.py actions rollback --run-id {record['run_id']}` to cancel created Things tasks.")
    finish_run(args, root, record, touched, "personal-os: apply Things actions")
    print_json({"run_id": record["run_id"], "created": created})


def cmd_actions_rollback(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    state_path = things_state_path(root)
    state = load_json(state_path, [])
    targets = [row for row in state if row.get("run_id") == args.run_id and not row.get("rolled_back")]
    if not targets:
        print_json({"run_id": args.run_id, "rolled_back": [], "message": "No active Things tasks found for run."})
        return
    record = start_run(args, "actions rollback", root)
    touched: list[Path] = []
    rolled_back = []
    for row in targets:
        item_id = row.get("id")
        if item_id:
            try:
                call_things(["cancel", str(item_id)])
                row["rollback_verification"] = verify_things_item(str(item_id))
                if row["rollback_verification"]:
                    row["rollback_effective_status"] = things_effective_status(row["rollback_verification"][0])
                row["rolled_back"] = True
                row["rolled_back_at"] = utc_now().isoformat().replace("+00:00", "Z")
                rolled_back.append(row)
            except SystemExit as exc:
                record["errors"].append(str(exc))
    write_json(state_path, state, touched)
    record["things_rolled_back"] = rolled_back
    record["rollback"].append("Rollback command only cancels tasks created by the target Personal OS run.")
    finish_run(args, root, record, touched, "personal-os: rollback Things actions")
    print_json({"run_id": record["run_id"], "rolled_back": rolled_back, "errors": record["errors"]})


def resolve_personal_path(root: Path, value: str) -> Path:
    path = Path(value).expanduser()
    if not path.is_absolute():
        path = root / path
    return ensure_under_root(root, path)


def iso_week_id(day: dt.date) -> str:
    year, week, _ = day.isocalendar()
    return f"{year}-W{week:02d}"


def parse_local_date(value: Any) -> dt.date | None:
    if value in (None, ""):
        return None
    text = str(value)
    if len(text) < 10:
        return None
    try:
        if "T" in text:
            normalized = text.replace("Z", "+00:00")
            parsed = dt.datetime.fromisoformat(normalized)
            if parsed.tzinfo is not None:
                return parsed.astimezone().date()
            return parsed.date()
        parsed_date = dt.date.fromisoformat(text[:10])
    except ValueError:
        return None
    if parsed_date.year >= 3999:
        return None
    return parsed_date


def display_date(value: Any) -> str:
    parsed = parse_local_date(value)
    return parsed.isoformat() if parsed else ""


def split_tag_names(value: Any) -> list[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [part.strip() for part in str(value or "").split(",") if part.strip()]


def read_things_snapshot_for_time() -> tuple[dict[str, Any], list[str]]:
    try:
        data = call_things(["snapshot-db", "--open-only"])
    except SystemExit as exc:
        primary_error = str(exc)
        try:
            data = call_things(["snapshot", "--open-only"])
            if isinstance(data, dict):
                data.setdefault("todos", [])
                data.setdefault("projects", [])
                return data, [f"Things read-only SQLite snapshot failed; used AppleScript fallback. Error: {primary_error}"]
        except SystemExit as fallback_exc:
            return {"todos": [], "projects": [], "areas": [], "tags": []}, [
                f"Things read-only SQLite snapshot failed: {primary_error}",
                f"Things AppleScript fallback failed: {fallback_exc}",
            ]
    if not isinstance(data, dict):
        return {"todos": [], "projects": [], "areas": [], "tags": []}, ["Things snapshot did not return a JSON object."]
    data.setdefault("todos", [])
    data.setdefault("projects", [])
    return data, []


def verify_agent_tasks_project_for_time(root: Path) -> tuple[dict[str, Any] | None, list[str]]:
    warnings: list[str] = []
    try:
        return verify_agent_tasks_project(root), warnings
    except SystemExit as exc:
        warnings.append(f"Things AppleScript project verification failed; trying read-only SQLite fallback. Error: {exc}")
    routing = load_things_routing(root)
    data = call_things(["snapshot-db", "--open-only"])
    for project in data.get("projects", []):
        if project.get("id") == routing["project_id"] and project.get("name") == routing["project_name"]:
            return project, warnings
    raise SystemExit(f"Things project not found by AppleScript or read-only fallback: {routing['project_name']}")


def project_open_counts(snapshot: dict[str, Any]) -> dict[str, int]:
    counts: dict[str, int] = {}
    for todo in snapshot.get("todos", []):
        if not isinstance(todo, dict):
            continue
        project = str(todo.get("projectName") or "(No Project)")
        counts[project] = counts.get(project, 0) + 1
    return counts


def things_pressure(snapshot: dict[str, Any], day: dt.date, horizon_end: dt.date) -> dict[str, int]:
    pressure = {
        "open": 0,
        "overdue_scheduled": 0,
        "scheduled_today": 0,
        "scheduled_horizon": 0,
        "deadline_overdue": 0,
        "deadline_today": 0,
        "deadline_horizon": 0,
        "stale_scheduled": 0,
        "stale_deadline": 0,
        "agent_tasks": 0,
        "personal_os": 0,
        "unscheduled": 0,
    }
    for todo in snapshot.get("todos", []):
        if not isinstance(todo, dict):
            continue
        pressure["open"] += 1
        when = parse_local_date(todo.get("when"))
        deadline = parse_local_date(todo.get("deadline"))
        tags = split_tag_names(todo.get("tagNames"))
        if when:
            if when < day:
                pressure["overdue_scheduled"] += 1
                if (day - when).days > 30:
                    pressure["stale_scheduled"] += 1
            elif when == day:
                pressure["scheduled_today"] += 1
            elif when <= horizon_end:
                pressure["scheduled_horizon"] += 1
        if deadline:
            if deadline < day:
                pressure["deadline_overdue"] += 1
                if (day - deadline).days > 30:
                    pressure["stale_deadline"] += 1
            elif deadline == day:
                pressure["deadline_today"] += 1
            elif deadline <= horizon_end:
                pressure["deadline_horizon"] += 1
        if not when and not deadline:
            pressure["unscheduled"] += 1
        if todo.get("projectName") == THINGS_PROJECT_NAME:
            pressure["agent_tasks"] += 1
        if THINGS_TAG in tags:
            pressure["personal_os"] += 1
    return pressure


def score_thing(todo: dict[str, Any], day: dt.date, horizon_end: dt.date, counts: dict[str, int]) -> tuple[int, list[str]]:
    score = 0
    reasons: list[str] = []
    when = parse_local_date(todo.get("when"))
    deadline = parse_local_date(todo.get("deadline"))
    tags = split_tag_names(todo.get("tagNames"))
    project = str(todo.get("projectName") or "(No Project)")
    if when:
        days = (when - day).days
        if days < 0:
            overdue_days = abs(days)
            if overdue_days > 30:
                score += 260
                reasons.append(f"stale scheduled date overdue by {overdue_days}d")
            elif overdue_days > 14:
                score += 650
                reasons.append(f"scheduled overdue by {overdue_days}d")
            else:
                score += 1000 + overdue_days
                reasons.append(f"scheduled overdue by {overdue_days}d")
        elif days == 0:
            score += 900
            reasons.append("scheduled today")
        elif when <= horizon_end:
            score += 720 - min(days * 20, 300)
            reasons.append(f"scheduled in {days}d")
    if deadline:
        days = (deadline - day).days
        if days < 0:
            overdue_days = abs(days)
            if overdue_days > 30:
                score += 300
                reasons.append(f"stale deadline overdue by {overdue_days}d")
            elif overdue_days > 14:
                score += 620
                reasons.append(f"deadline overdue by {overdue_days}d")
            else:
                score += 950 + overdue_days
                reasons.append(f"deadline overdue by {overdue_days}d")
        elif days == 0:
            score += 860
            reasons.append("deadline today")
        elif deadline <= horizon_end:
            score += 680 - min(days * 25, 350)
            reasons.append(f"deadline in {days}d")
    if project == THINGS_PROJECT_NAME:
        score += 520
        reasons.append("agent task queue")
    if THINGS_TAG in tags:
        score += 420
        reasons.append("Personal OS follow-up")
    if score == 0 and project != "(No Project)":
        score += min(counts.get(project, 0) * 4, 80)
        reasons.append("project backlog momentum")
    if score == 0:
        score = 10
        reasons.append("unscheduled backlog")
    return score, reasons


def ranked_things(snapshot: dict[str, Any], day: dt.date, horizon_days: int) -> list[dict[str, Any]]:
    horizon_end = day + dt.timedelta(days=horizon_days)
    counts = project_open_counts(snapshot)
    ranked: list[dict[str, Any]] = []
    for todo in snapshot.get("todos", []):
        if not isinstance(todo, dict):
            continue
        score, reasons = score_thing(todo, day, horizon_end, counts)
        ranked.append(
            {
                "id": todo.get("id", ""),
                "title": todo.get("name") or todo.get("title") or "Untitled",
                "project": todo.get("projectName") or "",
                "area": todo.get("areaName") or "",
                "tags": split_tag_names(todo.get("tagNames")),
                "when": display_date(todo.get("when")),
                "deadline": display_date(todo.get("deadline")),
                "score": score,
                "reasons": reasons,
            }
        )
    ranked.sort(key=lambda item: (-int(item["score"]), str(item["title"]).casefold()))
    return ranked


def calendar_context(day: dt.date, horizon_days: int) -> dict[str, Any]:
    if not CALENDAR_SCRIPT.exists():
        return {"events": [], "warnings": [f"Calendar helper not found: {CALENDAR_SCRIPT}"]}
    end = day + dt.timedelta(days=horizon_days + 1)
    proc = run_command([sys.executable, str(CALENDAR_SCRIPT), "agenda", "--from", day.isoformat(), "--to", end.isoformat()])
    if proc.returncode != 0:
        return {"events": [], "warnings": [proc.stderr.strip() or proc.stdout.strip() or "Calendar agenda failed."]}
    try:
        data = json.loads(proc.stdout or "[]")
    except json.JSONDecodeError:
        return {"events": [], "warnings": ["Calendar agenda returned invalid JSON."]}
    if not isinstance(data, list):
        return {"events": [], "warnings": ["Calendar agenda returned an unexpected shape."]}
    events = []
    for event in data:
        if isinstance(event, dict):
            events.append(
                {
                    "title": event.get("title", "Untitled"),
                    "start": event.get("start", ""),
                    "end": event.get("end", ""),
                    "calendar": event.get("calendarName", ""),
                    "all_day": bool(event.get("allDay")),
                    "location": event.get("location", ""),
                }
            )
    return {"events": events, "warnings": []}


def read_review_needed_items(root: Path, limit: int = 20) -> list[str]:
    path = root / "_views" / "review-needed.md"
    if not path.exists():
        return []
    items: list[str] = []
    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        clean = line.strip()
        if clean.startswith("- ") and not clean.startswith("- None"):
            items.append(clean[2:])
        if len(items) >= limit:
            break
    return items


def wiki_experiment_status() -> dict[str, Any]:
    candidates = [
        LLM_WIKI_ROOT / "_views" / "experiment-scan.md",
        LLM_WIKI_ROOT / "_views" / "agent-experiments.md",
        LLM_WIKI_ROOT / "_views" / "recent-sources.md",
    ]
    existing = [path for path in candidates if path.exists()]
    if not existing:
        return {"available": False, "views": [], "note": "No local wiki experiment scan view found."}
    views = []
    for path in existing:
        try:
            rel = str(path.relative_to(LLM_WIKI_ROOT))
        except ValueError:
            rel = str(path)
        views.append({"path": str(path), "wiki_relative": rel})
    return {"available": True, "views": views, "note": "External wiki experiment queues are separate from Things until promoted."}


def time_agent_candidates(data: dict[str, Any]) -> list[dict[str, str]]:
    candidates: list[dict[str, str]] = []
    pressure = data.get("things_pressure", {})
    review_items = data.get("personal_os_review", [])
    warnings = data.get("warnings", [])
    if int(pressure.get("agent_tasks", 0) or 0) > 0:
        candidates.append(
            {
                "title": f"Work through Agent Tasks queue ({pressure.get('agent_tasks')} open)",
                "notes": "Generated from Personal OS time planning because the dedicated agent queue has open tasks.",
                "source_key": "time:agent-tasks-open",
            }
        )
    stale_count = int(pressure.get("stale_scheduled", 0) or 0) + int(pressure.get("stale_deadline", 0) or 0)
    if stale_count > 0:
        candidates.append(
            {
                "title": f"Review stale dated Things items ({stale_count} date signals)",
                "notes": "Generated from Personal OS time planning because old scheduled/deadline dates may no longer represent real priority.",
                "source_key": "time:stale-dated-things",
            }
        )
    if review_items:
        candidates.append(
            {
                "title": f"Review Personal OS review-needed queue ({len(review_items)} visible items)",
                "notes": "\n".join(f"- {item}" for item in review_items),
                "source_key": "time:personal-os-review-needed",
            }
        )
    if warnings:
        for index, warning in enumerate(warnings):
            if "Calendar" in warning or "calendar" in warning:
                candidates.append(
                    {
                        "title": "Restore Calendar read access for Personal OS time planning",
                        "notes": warning,
                        "source_key": f"time:calendar-warning:{index}",
                    }
                )
                break
    return candidates


def build_time_daily_data(root: Path, day: dt.date, horizon_days: int, limit: int) -> dict[str, Any]:
    snapshot, things_warnings = read_things_snapshot_for_time()
    horizon_end = day + dt.timedelta(days=horizon_days)
    calendar = calendar_context(day, horizon_days)
    ranked = ranked_things(snapshot, day, horizon_days)
    pressure = things_pressure(snapshot, day, horizon_end)
    review_items = read_review_needed_items(root)
    wiki_status = wiki_experiment_status()
    data = {
        "type": "personal_os_time_daily",
        "date": day.isoformat(),
        "generated_at": utc_now().isoformat().replace("+00:00", "Z"),
        "horizon_days": horizon_days,
        "limit": limit,
        "warnings": [*things_warnings, *calendar.get("warnings", [])],
        "things_pressure": pressure,
        "ranked_next_actions": ranked[:limit],
        "calendar_events": calendar.get("events", []),
        "personal_os_review": review_items,
        "wiki_experiments": wiki_status,
    }
    data["agent_task_candidates"] = time_agent_candidates(data)
    return data


def build_time_weekly_data(root: Path, day: dt.date, horizon_days: int, limit: int) -> dict[str, Any]:
    snapshot, things_warnings = read_things_snapshot_for_time()
    horizon_end = day + dt.timedelta(days=horizon_days)
    ranked = ranked_things(snapshot, day, horizon_days)
    counts = project_open_counts(snapshot)
    scheduled_by_project: dict[str, int] = {}
    deadline_by_project: dict[str, int] = {}
    for todo in snapshot.get("todos", []):
        if not isinstance(todo, dict):
            continue
        project = str(todo.get("projectName") or "(No Project)")
        when = parse_local_date(todo.get("when"))
        deadline = parse_local_date(todo.get("deadline"))
        if when and when <= horizon_end:
            scheduled_by_project[project] = scheduled_by_project.get(project, 0) + 1
        if deadline and deadline <= horizon_end:
            deadline_by_project[project] = deadline_by_project.get(project, 0) + 1
    project_rows = []
    for project, count in sorted(counts.items(), key=lambda row: (-row[1], row[0].casefold())):
        project_rows.append(
            {
                "project": project,
                "open": count,
                "scheduled_horizon": scheduled_by_project.get(project, 0),
                "deadline_horizon": deadline_by_project.get(project, 0),
            }
        )
    data = {
        "type": "personal_os_time_weekly",
        "week": iso_week_id(day),
        "date": day.isoformat(),
        "generated_at": utc_now().isoformat().replace("+00:00", "Z"),
        "horizon_days": horizon_days,
        "warnings": things_warnings,
        "things_pressure": things_pressure(snapshot, day, horizon_end),
        "project_pressure": project_rows[:limit],
        "ranked_next_actions": ranked[:limit],
        "personal_os_review": read_review_needed_items(root),
        "wiki_experiments": wiki_experiment_status(),
    }
    data["agent_task_candidates"] = time_agent_candidates(data)
    return data


def render_time_daily(data: dict[str, Any]) -> str:
    lines = [
        frontmatter(
            {
                "type": data["type"],
                "date": data["date"],
                "updated": data["generated_at"],
                "horizon_days": data["horizon_days"],
            }
        ).rstrip(),
        "",
        f"# Time Daily {data['date']}",
        "",
        "## Summary",
        "",
    ]
    pressure = data["things_pressure"]
    lines.extend(
        [
            f"- Open Things tasks: {pressure.get('open', 0)}",
            f"- Scheduled overdue/today/horizon: {pressure.get('overdue_scheduled', 0)}/{pressure.get('scheduled_today', 0)}/{pressure.get('scheduled_horizon', 0)}",
            f"- Deadline overdue/today/horizon: {pressure.get('deadline_overdue', 0)}/{pressure.get('deadline_today', 0)}/{pressure.get('deadline_horizon', 0)}",
            f"- Stale dated signals: scheduled {pressure.get('stale_scheduled', 0)}, deadline {pressure.get('stale_deadline', 0)}",
            f"- Agent Tasks open: {pressure.get('agent_tasks', 0)}",
            f"- Personal OS follow-ups open: {pressure.get('personal_os', 0)}",
            "",
        ]
    )
    if data.get("warnings"):
        lines.extend(["## Warnings", ""])
        for warning in data["warnings"]:
            lines.append(f"- {warning}")
        lines.append("")
    lines.extend(["## Calendar Context", ""])
    events = data.get("calendar_events", [])
    if events:
        for event in events[:40]:
            prefix = "all day" if event.get("all_day") else f"{event.get('start', '')} - {event.get('end', '')}"
            suffix = f" ({event.get('calendar')})" if event.get("calendar") else ""
            lines.append(f"- {prefix}: {event.get('title', 'Untitled')}{suffix}")
    else:
        lines.append("- No events found or Calendar read was unavailable.")
    lines.extend(["", "## Ranked Next Actions", ""])
    for index, item in enumerate(data.get("ranked_next_actions", []), start=1):
        where = item.get("project") or item.get("area") or "Inbox"
        schedule = []
        if item.get("when"):
            schedule.append(f"when {item['when']}")
        if item.get("deadline"):
            schedule.append(f"deadline {item['deadline']}")
        schedule_text = "; ".join(schedule) or "unscheduled"
        reasons = ", ".join(item.get("reasons", []))
        lines.append(f"{index}. {item.get('title')} - {where}; {schedule_text}; {reasons}; id `{item.get('id')}`")
    if not data.get("ranked_next_actions"):
        lines.append("- No Things actions available.")
    lines.extend(["", "## Personal OS Review Pressure", ""])
    for item in data.get("personal_os_review", [])[:20]:
        lines.append(f"- {item}")
    if not data.get("personal_os_review"):
        lines.append("- None visible.")
    lines.extend(["", "## Separate External Experiment Queue", ""])
    wiki = data.get("wiki_experiments", {})
    lines.append(f"- {wiki.get('note', '')}")
    for view in wiki.get("views", []):
        lines.append(f"- `{view.get('wiki_relative')}`")
    lines.extend(["", "## Agent Task Candidates", ""])
    if data.get("agent_task_candidates"):
        for candidate in data["agent_task_candidates"]:
            lines.append(f"- ACTION: {candidate['title']}")
    else:
        lines.append("- None.")
    lines.append("")
    return "\n".join(lines)


def render_time_weekly(data: dict[str, Any]) -> str:
    lines = [
        frontmatter(
            {
                "type": data["type"],
                "week": data["week"],
                "date": data["date"],
                "updated": data["generated_at"],
                "horizon_days": data["horizon_days"],
            }
        ).rstrip(),
        "",
        f"# Time Weekly {data['week']}",
        "",
        "## Summary",
        "",
    ]
    pressure = data["things_pressure"]
    lines.extend(
        [
            f"- Open Things tasks: {pressure.get('open', 0)}",
            f"- Scheduled in horizon: {pressure.get('scheduled_horizon', 0)} plus today {pressure.get('scheduled_today', 0)} and overdue {pressure.get('overdue_scheduled', 0)}",
            f"- Deadlines in horizon: {pressure.get('deadline_horizon', 0)} plus today {pressure.get('deadline_today', 0)} and overdue {pressure.get('deadline_overdue', 0)}",
            f"- Stale dated signals: scheduled {pressure.get('stale_scheduled', 0)}, deadline {pressure.get('stale_deadline', 0)}",
            f"- Unscheduled backlog: {pressure.get('unscheduled', 0)}",
            "",
        ]
    )
    if data.get("warnings"):
        lines.extend(["## Warnings", ""])
        for warning in data["warnings"]:
            lines.append(f"- {warning}")
        lines.append("")
    lines.extend(["## Project Pressure", ""])
    for row in data.get("project_pressure", []):
        lines.append(
            f"- {row['project']}: {row['open']} open; {row['scheduled_horizon']} scheduled; {row['deadline_horizon']} deadlines"
        )
    if not data.get("project_pressure"):
        lines.append("- No open project pressure.")
    lines.extend(["", "## Top Next Actions", ""])
    for index, item in enumerate(data.get("ranked_next_actions", []), start=1):
        where = item.get("project") or item.get("area") or "Inbox"
        reasons = ", ".join(item.get("reasons", []))
        lines.append(f"{index}. {item.get('title')} - {where}; {reasons}; id `{item.get('id')}`")
    if not data.get("ranked_next_actions"):
        lines.append("- No Things actions available.")
    lines.extend(["", "## Agent Task Candidates", ""])
    if data.get("agent_task_candidates"):
        for candidate in data["agent_task_candidates"]:
            lines.append(f"- ACTION: {candidate['title']}")
    else:
        lines.append("- None.")
    lines.append("")
    return "\n".join(lines)


def action_plan_from_time_report(root: Path, source: Path) -> dict[str, Any]:
    text = source.read_text(encoding="utf-8")
    tasks = []
    source_rel = relative(root, source)
    for title in explicit_actions_from_text(text):
        tasks.append(
            {
                "title": title,
                "notes": f"Created from Personal OS time report: {source_rel}",
                "source": source_rel,
                "source_key": f"time:{source_rel}:{slugify(title)}",
                "tag": THINGS_TAG,
            }
        )
    return {
        "type": "personal_os_action_plan",
        "created": utc_now().isoformat().replace("+00:00", "Z"),
        "source_reflection": source_rel,
        "source_report": source_rel,
        "tasks": tasks,
    }


def cmd_time_daily(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    day = parse_day(args.date)
    data = build_time_daily_data(root, day, args.horizon_days, args.limit)
    rendered = render_time_daily(data)
    if not args.write:
        print(rendered)
        return
    record = start_run(args, "time daily", root)
    touched: list[Path] = []
    report_path = root / "_reports" / "time" / f"{day.isoformat()}-daily.md"
    write_text(root / "_views" / "time-daily.md", rendered, touched)
    write_text(report_path, rendered, touched)
    write_json(root / "_state" / "time" / "daily.json", data, touched)
    write_text(root / "_views" / "agent-task-candidates.md", render_agent_task_candidates_view(data, "time-daily"), touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for the time daily commit, or remove {relative(root, report_path)}.")
    finish_run(args, root, record, touched, f"personal-os: time daily {day.isoformat()}")
    print_json({"run_id": record["run_id"], "report": str(report_path), "view": str(root / "_views" / "time-daily.md")})


def cmd_time_weekly(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    day = parse_day(args.date)
    data = build_time_weekly_data(root, day, args.horizon_days, args.limit)
    rendered = render_time_weekly(data)
    if not args.write:
        print(rendered)
        return
    record = start_run(args, "time weekly", root)
    touched: list[Path] = []
    week = iso_week_id(day)
    report_path = root / "_reports" / "time" / f"{week}-weekly.md"
    write_text(root / "_views" / "time-weekly.md", rendered, touched)
    write_text(report_path, rendered, touched)
    write_json(root / "_state" / "time" / "weekly.json", data, touched)
    write_text(root / "_views" / "agent-task-candidates.md", render_agent_task_candidates_view(data, "time-weekly"), touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append(f"Use git revert for the time weekly commit, or remove {relative(root, report_path)}.")
    finish_run(args, root, record, touched, f"personal-os: time weekly {week}")
    print_json({"run_id": record["run_id"], "report": str(report_path), "view": str(root / "_views" / "time-weekly.md")})


def render_agent_task_candidates_view(data: dict[str, Any], source_view: str) -> str:
    lines = [
        frontmatter(
            {
                "type": "view",
                "view": "agent-task-candidates",
                "source_view": source_view,
                "updated": data.get("generated_at", utc_now().isoformat().replace("+00:00", "Z")),
            }
        ).rstrip(),
        "",
        "# Agent Task Candidates",
        "",
        "These are explicit candidate actions generated from Personal OS time planning. Apply them with `personal.py time agent-plan` and `personal.py time agent-apply` only when they should become Things tasks.",
        "",
    ]
    candidates = data.get("agent_task_candidates", [])
    if not candidates:
        lines.append("- None.")
    for candidate in candidates:
        lines.append(f"- ACTION: {candidate['title']}")
        if candidate.get("notes"):
            lines.append(f"  - Notes: {candidate['notes'].splitlines()[0]}")
    lines.append("")
    return "\n".join(lines)


def cmd_time_agent_plan(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    source = resolve_personal_path(root, args.from_report)
    plan = action_plan_from_time_report(root, source)
    if not args.write:
        print_json(plan)
        return
    record = start_run(args, "time agent-plan", root)
    touched: list[Path] = []
    plan_path = root / "_views" / f"action-plan-{record['run_id']}.json"
    write_json(plan_path, plan, touched)
    record["rollback"].append(f"Remove {relative(root, plan_path)} or use git revert.")
    finish_run(args, root, record, touched, "personal-os: time action plan")
    print_json({"run_id": record["run_id"], "plan": str(plan_path), "task_count": len(plan["tasks"])})


def cmd_time_agent_apply(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    plan_path = resolve_personal_path(root, args.plan)
    plan = load_action_plan(root, plan_path)
    record = start_run(args, "time agent-apply", root)
    touched: list[Path] = []
    state = load_json(things_state_path(root), [])
    created = []
    verify_agent_tasks_project(root)
    for task in plan.get("tasks", []):
        title = str(task.get("title", "")).strip()
        if not title:
            continue
        notes = f"{task.get('notes', '')}\n\nPersonal OS run: {record['run_id']}"
        source = str(task.get("source", plan.get("source_report", plan.get("source_reflection", ""))))
        source_key = str(task.get("source_key") or f"time:{source}:{slugify(title)}")
        created_row = create_personal_os_things_task(
            root,
            title=title,
            notes=notes,
            source=source,
            source_key=source_key,
            run_id_value=record["run_id"],
            state=state,
        )
        if created_row:
            created_row["plan"] = relative(root, plan_path)
            created.append(created_row)
    write_json(things_state_path(root), state, touched)
    record["things_created"] = created
    record["rollback"].append(f"Run `personal.py actions rollback --run-id {record['run_id']}` to cancel created Things tasks.")
    finish_run(args, root, record, touched, "personal-os: apply time agent actions")
    print_json({"run_id": record["run_id"], "created": created})


def verify_time_layer(root: Path) -> list[str]:
    errors: list[str] = []
    for name in ("_state/time", "_reports/time"):
        if not (root / name).is_dir():
            errors.append(f"Missing time directory: {name}")
    for path in (root / "_state" / "time").glob("*.json"):
        try:
            json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            errors.append(f"Invalid time state JSON: {relative(root, path)}")
    if not THINGS_SCRIPT.exists():
        errors.append(f"Things helper missing: {THINGS_SCRIPT}")
    if not CALENDAR_SCRIPT.exists():
        errors.append(f"Calendar helper missing: {CALENDAR_SCRIPT}")
    return errors


def cmd_time_verify(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    errors = verify_time_layer(root)
    warnings: list[str] = []
    try:
        _, project_warnings = verify_agent_tasks_project_for_time(root)
        warnings.extend(project_warnings)
    except SystemExit as exc:
        errors.append(str(exc))
    calendar = calendar_context(dt.date.today(), 0)
    warnings.extend(calendar.get("warnings", []))
    print_json(
        {
            "root": str(root),
            "ok": not errors,
            "errors": errors,
            "warnings": warnings,
            "latest_daily_view": str(root / "_views" / "time-daily.md") if (root / "_views" / "time-daily.md").exists() else "",
            "latest_weekly_view": str(root / "_views" / "time-weekly.md") if (root / "_views" / "time-weekly.md").exists() else "",
        }
    )
    if errors:
        raise SystemExit(1)


def garden_task_candidates(root: Path, manifests: list[dict[str, Any]]) -> list[dict[str, str]]:
    needs_ocr = [manifest for manifest in manifests if manifest.get("extraction_status") == "needs_ocr"]
    extraction_failures = [manifest for manifest in manifests if manifest.get("extraction_status") in {"failed", "unsupported"}]
    pending_agent = [manifest for manifest in manifests if manifest.get("summary_status") == "pending_agent"]
    duplicates = [manifest for manifest in manifests if manifest.get("duplicate_of")]
    candidates: list[dict[str, str]] = []
    if needs_ocr:
        candidates.append(
            {
                "title": f"Review Personal OS OCR failures ({len(needs_ocr)} files)",
                "notes": "\n".join(f"- {item.get('title')} ({item.get('file_id')})" for item in needs_ocr),
                "source": "_views/review-needed.md",
                "source_key": "garden:needs-ocr",
            }
        )
    if extraction_failures:
        candidates.append(
            {
                "title": f"Review Personal OS extraction failures ({len(extraction_failures)} files)",
                "notes": "\n".join(f"- {item.get('title')} ({item.get('file_id')}): {item.get('extraction_error', '')}" for item in extraction_failures),
                "source": "_views/review-needed.md",
                "source_key": "garden:extraction-failures",
            }
        )
    if pending_agent:
        candidates.append(
            {
                "title": f"Ask Codex to synthesize pending Personal OS file summaries ({len(pending_agent)} files)",
                "notes": "\n".join(f"- {item.get('title')} ({item.get('file_id')})" for item in pending_agent),
                "source": "_views/review-needed.md",
                "source_key": "garden:pending-agent-summaries",
            }
        )
    if duplicates:
        candidates.append(
            {
                "title": f"Decide what to do with duplicate Personal OS files ({len(duplicates)} files)",
                "notes": "\n".join(f"- {item.get('title')} ({item.get('file_id')}) duplicates {item.get('duplicate_of')}" for item in duplicates),
                "source": "_views/review-needed.md",
                "source_key": "garden:duplicates",
            }
        )
    return candidates


def garden_report(
    root: Path,
    day: dt.date,
    *,
    rechecked: list[dict[str, Any]],
    extracted: list[dict[str, Any]],
    summarized: list[dict[str, Any]],
    tasks: list[dict[str, Any]],
    errors: list[str],
) -> str:
    manifests = load_file_manifests(root)
    review_rows = review_needed_rows(manifests)
    lines = [
        frontmatter(
            {
                "type": "garden_report",
                "date": day.isoformat(),
                "updated": utc_now().isoformat().replace("+00:00", "Z"),
                "file_count": len(manifests),
                "review_count": len(review_rows),
            }
        ).rstrip(),
        "",
        f"# Personal OS Garden Report {day.isoformat()}",
        "",
        "## Summary",
        "",
        f"- Files tracked: {len(manifests)}",
        f"- Rechecked: {len(rechecked)}",
        f"- Extracted: {len(extracted)}",
        f"- Intake summaries/outlines updated: {len(summarized)}",
        f"- Things tasks created: {len(tasks)}",
        f"- Errors: {len(errors)}",
        "",
        "## Review Needed",
        "",
    ]
    if review_rows:
        for row in review_rows[:80]:
            manifest = row["manifest"]
            lines.append(f"- {manifest.get('title')} (`{manifest.get('file_id')}`): {', '.join(row['reasons'])}")
    else:
        lines.append("- None.")
    lines.extend(["", "## Things Tasks Created", ""])
    if tasks:
        for task in tasks:
            lines.append(f"- {task.get('title')} (`{task.get('id', '')}`)")
    else:
        lines.append("- None.")
    if errors:
        lines.extend(["", "## Errors", ""])
        for error in errors:
            lines.append(f"- {error}")
    lines.append("")
    return "\n".join(lines)


def cmd_garden_daily(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    day = parse_day(args.date)
    manifests = load_file_manifests(root)
    pending_extract = [m for m in manifests if m.get("extraction_status") in ("pending", None)]
    pending_summary = [m for m in manifests if m.get("summary_status") in ("pending", None)]
    if not args.write:
        print_json(
            {
                "date": day.isoformat(),
                "file_count": len(manifests),
                "pending_extract_count": len(pending_extract),
                "pending_summary_count": len(pending_summary),
                "write_required": True,
            }
        )
        return
    record = start_run(args, "garden daily", root)
    touched: list[Path] = []
    errors: list[str] = []
    rechecked = recheck_manifests(root, touched, write=True)
    extracted = []
    summarized = []
    for manifest in load_file_manifests(root):
        try:
            if manifest.get("extraction_status") in ("pending", None) or manifest.get("source_changed"):
                manifest = extract_manifest(root, manifest, touched)
                extracted.append(manifest)
            if manifest.get("summary_status") in ("pending", None) or manifest.get("source_changed"):
                manifest = summarize_manifest(root, manifest, touched)
                summarized.append(manifest)
        except Exception as exc:  # Keep the garden run useful even if one file fails.
            errors.append(f"{manifest.get('file_id', '<unknown>')}: {exc}")
    rebuild_file_views(root, touched)
    manifests = load_file_manifests(root)
    state = load_json(things_state_path(root), [])
    if not isinstance(state, list):
        state = []
    tasks_created = []
    candidates = garden_task_candidates(root, manifests)
    if candidates:
        try:
            verify_agent_tasks_project(root)
            for candidate in candidates:
                created = create_personal_os_things_task(
                    root,
                    title=candidate["title"],
                    notes=candidate["notes"],
                    source=candidate["source"],
                    source_key=candidate["source_key"],
                    run_id_value=record["run_id"],
                    state=state,
                )
                if created:
                    tasks_created.append(created)
        except SystemExit as exc:
            errors.append(str(exc))
    if candidates or tasks_created:
        write_json(things_state_path(root), state, touched)
    report_text = garden_report(root, day, rechecked=rechecked, extracted=extracted, summarized=summarized, tasks=tasks_created, errors=errors)
    report_path = root / "_reports" / "garden" / f"{day.isoformat()}.md"
    write_text(report_path, report_text, touched)
    write_text(root / "_views" / "garden-report-latest.md", report_text, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["things_created"] = tasks_created
    record["errors"] = errors
    record["rollback"].append(f"Use git revert for the garden commit, and run `personal.py actions rollback --run-id {record['run_id']}` for created Things tasks.")
    finish_run(args, root, record, touched, f"personal-os: garden daily {day.isoformat()}")
    print_json(
        {
            "run_id": record["run_id"],
            "report": str(report_path),
            "rechecked_count": len(rechecked),
            "extracted_count": len(extracted),
            "summarized_count": len(summarized),
            "things_created_count": len(tasks_created),
            "errors": errors,
        }
    )


def cmd_butler_rebuild(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    if not args.write:
        print(butlers_book_view(root, load_file_manifests(root)))
        return
    record = start_run(args, "butler rebuild", root)
    touched: list[Path] = []
    rebuild_file_views(root, touched)
    write_text(root / "index.md", render_index(root), touched)
    record["rollback"].append("Use git revert for the butler rebuild commit.")
    finish_run(args, root, record, touched, "personal-os: rebuild butler views")
    print_json({"run_id": record["run_id"], "butlers_book": str(root / "_views" / "butlers-book.md")})


def cmd_automation_install(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    record = start_run(args, "automation install", root)
    touched: list[Path] = []
    write_plist(LAUNCH_AGENT_PATH, launch_agent_plist(root, disabled=True))
    record["automation"].append({"action": "install", "path": str(LAUNCH_AGENT_PATH), "loaded": False})
    record["rollback"].append("Run `personal.py automation uninstall` to remove the LaunchAgent plist.")
    finish_run(args, root, record, touched, "personal-os: install automation")
    print_json({"run_id": record["run_id"], "plist": str(LAUNCH_AGENT_PATH), "loaded": False})


def launchctl(*args: str) -> dict[str, Any]:
    proc = run_command(["launchctl", *args])
    return {"args": list(args), "returncode": proc.returncode, "stdout": proc.stdout.strip(), "stderr": proc.stderr.strip()}


def cmd_automation_enable(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    if not LAUNCH_AGENT_PATH.exists():
        raise SystemExit("LaunchAgent is not installed. Run automation install first.")
    record = start_run(args, "automation enable", root)
    touched: list[Path] = []
    write_plist(LAUNCH_AGENT_PATH, launch_agent_plist(root, disabled=False))
    uid = str(os.getuid())
    results = [
        launchctl("bootout", f"gui/{uid}", str(LAUNCH_AGENT_PATH)),
        launchctl("bootstrap", f"gui/{uid}", str(LAUNCH_AGENT_PATH)),
        launchctl("enable", f"gui/{uid}/{LAUNCH_AGENT_LABEL}"),
    ]
    record["automation"].append({"action": "enable", "path": str(LAUNCH_AGENT_PATH), "launchctl": results})
    record["rollback"].append("Run `personal.py automation disable` to unload the LaunchAgent.")
    finish_run(args, root, record, touched, "personal-os: enable automation")
    print_json({"run_id": record["run_id"], "launchctl": results})


def cmd_automation_disable(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    record = start_run(args, "automation disable", root)
    touched: list[Path] = []
    uid = str(os.getuid())
    results = [launchctl("bootout", f"gui/{uid}", str(LAUNCH_AGENT_PATH))]
    if LAUNCH_AGENT_PATH.exists():
        write_plist(LAUNCH_AGENT_PATH, launch_agent_plist(root, disabled=True))
    record["automation"].append({"action": "disable", "path": str(LAUNCH_AGENT_PATH), "launchctl": results})
    record["rollback"].append("Run `personal.py automation enable` to load the LaunchAgent again.")
    finish_run(args, root, record, touched, "personal-os: disable automation")
    print_json({"run_id": record["run_id"], "launchctl": results})


def cmd_automation_uninstall(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    record = start_run(args, "automation uninstall", root)
    touched: list[Path] = []
    uid = str(os.getuid())
    results = [launchctl("bootout", f"gui/{uid}", str(LAUNCH_AGENT_PATH))]
    removed = False
    if LAUNCH_AGENT_PATH.exists():
        LAUNCH_AGENT_PATH.unlink()
        removed = True
    record["automation"].append({"action": "uninstall", "path": str(LAUNCH_AGENT_PATH), "removed": removed, "launchctl": results})
    record["rollback"].append("Run `personal.py automation install` to recreate the disabled LaunchAgent plist.")
    finish_run(args, root, record, touched, "personal-os: uninstall automation")
    print_json({"run_id": record["run_id"], "removed": removed, "launchctl": results})


def cmd_scan_recent(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    print_json({"root": str(root), "items": collect_notes(root)[: args.limit]})


def cmd_index(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    ensure_root_exists(root)
    rendered = render_index(root)
    if not args.write:
        print(rendered)
        return
    record = start_run(args, "index", root)
    touched: list[Path] = []
    write_text(root / "index.md", rendered, touched)
    record["rollback"].append("Use git revert for the index update commit.")
    finish_run(args, root, record, touched, "personal-os: update index")
    print_json({"run_id": record["run_id"], "index": str(root / "index.md")})


def verify_frontmatter(root: Path) -> list[str]:
    errors: list[str] = []
    for path in root.glob("**/*.md"):
        if ".git" in path.parts:
            continue
        try:
            rel_parts = path.relative_to(root).parts
        except ValueError:
            rel_parts = path.parts
        if len(rel_parts) >= 2 and rel_parts[0] == "files" and rel_parts[1] == "originals":
            continue
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            errors.append(f"Non-UTF8 Markdown file: {relative(root, path)}")
            continue
        if not text.startswith("---\n"):
            errors.append(f"Missing frontmatter: {relative(root, path)}")
    return errors


def verify_run_logs(root: Path, run_logs: list[Path]) -> tuple[list[str], dict[str, Any]]:
    errors: list[str] = []
    latest_daily: dict[str, Any] | None = None
    for path in run_logs:
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
        except json.JSONDecodeError:
            errors.append(f"Invalid run log JSON: {relative(root, path)}")
            continue
        for field in ("run_id", "command", "started_at", "finished_at", "rollback"):
            if field not in data:
                errors.append(f"Run log missing {field}: {relative(root, path)}")
        for touched in data.get("files_touched", []):
            touched_path = root / touched
            if not touched_path.exists():
                errors.append(f"Run log touched missing file: {relative(root, path)} -> {touched}")
        if data.get("things_created") and data.get("command") not in {"actions apply", "garden daily", "time agent-apply"}:
            errors.append(f"Unexpected Things creation log: {relative(root, path)}")
        if data.get("command") == "reflect daily":
            if latest_daily is None or str(data.get("started_at", "")) > str(latest_daily.get("started_at", "")):
                latest_daily = data
    return errors, latest_daily or {}


def launch_agent_state(root: Path) -> dict[str, Any]:
    template_path = root / "_state" / "launchagents" / f"{LAUNCH_AGENT_LABEL}.plist"
    state: dict[str, Any] = {
        "template_path": str(template_path),
        "template_exists": template_path.exists(),
        "installed_path": str(LAUNCH_AGENT_PATH),
        "installed": LAUNCH_AGENT_PATH.exists(),
        "installed_disabled": None,
        "launchctl_loaded": False,
        "launchctl_detail": "",
    }
    if LAUNCH_AGENT_PATH.exists():
        try:
            with LAUNCH_AGENT_PATH.open("rb") as handle:
                state["installed_disabled"] = bool(plistlib.load(handle).get("Disabled", False))
        except Exception as exc:  # plistlib can raise several parse errors.
            state["installed_error"] = str(exc)
    uid = str(os.getuid())
    result = launchctl("print", f"gui/{uid}/{LAUNCH_AGENT_LABEL}")
    state["launchctl_loaded"] = result["returncode"] == 0
    state["launchctl_detail"] = result["stdout"] or result["stderr"]
    return state


def verify_things_state(root: Path) -> list[str]:
    errors: list[str] = []
    state_path = things_state_path(root)
    if not state_path.exists():
        return errors
    try:
        data = json.loads(state_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return [f"Invalid Things state JSON: {relative(root, state_path)}"]
    if not isinstance(data, list):
        return [f"Things state must be a list: {relative(root, state_path)}"]
    for index, row in enumerate(data):
        if not isinstance(row, dict):
            errors.append(f"Things state row is not an object at index {index}")
            continue
        for field in ("id", "title", "source", "run_id", "created_at", "rolled_back"):
            if field not in row:
                errors.append(f"Things state row {index} missing {field}")
        if row.get("rolled_back") and row.get("rollback_effective_status") not in ("canceled", "completed"):
            errors.append(f"Things state row {index} marked rolled_back without canceled/completed verification")
    return errors


def cmd_verify(args: argparse.Namespace) -> None:
    root = root_from_args(args)
    errors: list[str] = []
    missing_dirs = [name for name in REQUIRED_DIRS if not (root / name).is_dir()]
    errors.extend(f"Missing directory: {name}" for name in missing_dirs)
    for name in ("AGENTS.md", "rules.md", "index.md"):
        if not (root / name).exists():
            errors.append(f"Missing file: {name}")
    errors.extend(verify_frontmatter(root) if root.exists() else ["Root missing"])
    git_status = None
    if (root / ".git").exists():
        proc = run_command(["git", "-C", str(root), "status", "--porcelain"])
        git_status = proc.stdout.strip()
    else:
        errors.append("Missing private git repo")
    run_logs = sorted((root / "_logs" / "runs").glob("*.json")) if root.exists() else []
    run_log_errors, latest_daily_run = verify_run_logs(root, run_logs)
    errors.extend(run_log_errors)
    errors.extend(verify_things_state(root) if root.exists() else [])
    errors.extend(verify_file_layer(root) if root.exists() else [])
    errors.extend(verify_time_layer(root) if root.exists() else [])
    automation = launch_agent_state(root) if root.exists() else {}
    if root.exists() and not automation.get("template_exists"):
        errors.append("Missing disabled LaunchAgent template in _state/launchagents")
    latest_daily_reflection = None
    if root.exists():
        reflection_files = sorted((root / "reflections" / "daily").glob("*.md"), reverse=True)
        latest_daily_reflection = relative(root, reflection_files[0]) if reflection_files else None
    print_json(
        {
            "root": str(root),
            "ok": not errors,
            "errors": errors,
            "git_head": git_head(root),
            "git_dirty": bool(git_status),
            "git_status": git_status,
            "automation": automation,
            "latest_daily_reflection": latest_daily_reflection,
            "latest_daily_run_id": latest_daily_run.get("run_id"),
            "latest_daily_run_at": latest_daily_run.get("started_at"),
            "run_log_count": len(run_logs),
        }
    )
    if errors:
        raise SystemExit(1)


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Maintain the private Personal OS memory layer.")
    parser.add_argument("--root", default=str(DEFAULT_ROOT), help="Personal OS root. Defaults to ~/.codex/personal-os.")
    parser.add_argument("--no-commit", action="store_true", help="Write files without committing the Personal OS git repo.")
    sub = parser.add_subparsers(dest="command", required=True)

    bootstrap = sub.add_parser("bootstrap", help="Create the Personal OS root.")
    bootstrap.set_defaults(func=cmd_bootstrap)

    verify = sub.add_parser("verify", help="Verify Personal OS structure and state.")
    verify.set_defaults(func=cmd_verify)

    capture = sub.add_parser("capture", help="Capture journal or episode text.")
    capture_sub = capture.add_subparsers(dest="capture_command", required=True)
    journal = capture_sub.add_parser("journal", help="Append a journal section for a day.")
    journal.add_argument("--title", required=True)
    journal.add_argument("--date", default="today")
    journal.add_argument("--stdin", action="store_true")
    journal.set_defaults(func=cmd_capture_journal)
    episode = capture_sub.add_parser("episode", help="Create an explicit episode note.")
    episode.add_argument("--source", choices=("codex", "chronicle", "manual"), required=True)
    episode.add_argument("--title", required=True)
    episode.add_argument("--date", default="today")
    episode.add_argument("--stdin", action="store_true")
    episode.set_defaults(func=cmd_capture_episode)

    reflect = sub.add_parser("reflect", help="Generate reflections.")
    reflect_sub = reflect.add_subparsers(dest="reflect_command", required=True)
    daily = reflect_sub.add_parser("daily", help="Generate a daily reflection.")
    daily.add_argument("--date", default="today")
    daily.add_argument("--write", action="store_true")
    daily.add_argument("--no-chronicle", action="store_true")
    daily.set_defaults(func=cmd_reflect_daily)

    profile = sub.add_parser("profile", help="Work with durable profile notes.")
    profile_sub = profile.add_subparsers(dest="profile_command", required=True)
    profile_update = profile_sub.add_parser("update", help="Apply explicit profile candidates from a reflection.")
    profile_update.add_argument("--from-reflection", required=True)
    profile_update.add_argument("--apply", action="store_true")
    profile_update.set_defaults(func=cmd_profile_update)

    actions = sub.add_parser("actions", help="Plan, apply, or rollback Things actions.")
    actions_sub = actions.add_subparsers(dest="actions_command", required=True)
    actions_plan = actions_sub.add_parser("plan", help="Extract explicit Things action candidates.")
    actions_plan.add_argument("--from-reflection", required=True)
    actions_plan.add_argument("--write", action="store_true")
    actions_plan.set_defaults(func=cmd_actions_plan)
    actions_apply = actions_sub.add_parser("apply", help="Create Things tasks from an action plan.")
    actions_apply.add_argument("--plan", required=True)
    actions_apply.set_defaults(func=cmd_actions_apply)
    actions_rollback = actions_sub.add_parser("rollback", help="Cancel Things tasks created by a Personal OS run.")
    actions_rollback.add_argument("--run-id", required=True)
    actions_rollback.set_defaults(func=cmd_actions_rollback)

    file_cmd = sub.add_parser("file", help="Add, inspect, extract, summarize, search, or verify Personal OS files.")
    file_sub = file_cmd.add_subparsers(dest="file_command", required=True)
    file_add = file_sub.add_parser("add", help="Add one file to Personal OS.")
    file_add.add_argument("path")
    file_add.add_argument("--kind", choices=FILE_KINDS, default="other")
    file_add.add_argument("--sensitivity", choices=SENSITIVITY_LABELS)
    file_add.add_argument("--title")
    file_add.add_argument("--notes", default="")
    copy_group = file_add.add_mutually_exclusive_group()
    copy_group.add_argument("--copy", dest="copy_mode", action="store_const", const="copy", help="Copy the raw file into Personal OS. Default.")
    copy_group.add_argument("--external", dest="copy_mode", action="store_const", const="external", help="Track the file without copying the raw file.")
    file_add.set_defaults(func=cmd_file_add, copy_mode="copy")

    file_add_folder = file_sub.add_parser("add-folder", help="Add files from a folder to Personal OS.")
    file_add_folder.add_argument("path")
    file_add_folder.add_argument("--kind", choices=FILE_KINDS, default="other")
    file_add_folder.add_argument("--sensitivity", choices=SENSITIVITY_LABELS)
    file_add_folder.add_argument("--notes", default="")
    file_add_folder.add_argument("--recursive", action="store_true")
    file_add_folder.add_argument("--dry-run", action="store_true", help="Preview only. This is the default without --write.")
    file_add_folder.add_argument("--write", action="store_true")
    file_add_folder.set_defaults(func=cmd_file_add_folder)

    file_inspect = file_sub.add_parser("inspect", help="Inspect one Personal OS file manifest.")
    file_inspect.add_argument("file_id")
    file_inspect.set_defaults(func=cmd_file_inspect)

    file_extract = file_sub.add_parser("extract", help="Extract text for one Personal OS file.")
    file_extract.add_argument("--file-id", required=True)
    file_extract.add_argument("--write", action="store_true")
    file_extract.set_defaults(func=cmd_file_extract)

    file_summarize = file_sub.add_parser("summarize", help="Write a deterministic intake summary for one Personal OS file.")
    file_summarize.add_argument("--file-id", required=True)
    file_summarize.add_argument("--write", action="store_true")
    file_summarize.set_defaults(func=cmd_file_summarize)

    file_search = file_sub.add_parser("search", help="Search Personal OS file manifests, summaries, and extracted text.")
    file_search.add_argument("query")
    file_search.add_argument("--limit", type=int, default=30)
    file_search.set_defaults(func=cmd_file_search)

    file_verify = file_sub.add_parser("verify", help="Verify Personal OS file layer.")
    file_verify.set_defaults(func=cmd_file_verify)

    file_recheck = file_sub.add_parser("recheck", help="Recheck source/copied file hashes.")
    file_recheck.add_argument("--write", action="store_true")
    file_recheck.set_defaults(func=cmd_file_recheck)

    time_cmd = sub.add_parser("time", help="Generate time and next-action planning surfaces.")
    time_sub = time_cmd.add_subparsers(dest="time_command", required=True)
    time_daily = time_sub.add_parser("daily", help="Generate a daily time and next-action plan.")
    time_daily.add_argument("--date", default="today")
    time_daily.add_argument("--horizon-days", type=int, default=7)
    time_daily.add_argument("--limit", type=int, default=10)
    time_daily.add_argument("--write", action="store_true")
    time_daily.set_defaults(func=cmd_time_daily)
    time_weekly = time_sub.add_parser("weekly", help="Generate a weekly project-pressure plan.")
    time_weekly.add_argument("--date", default="today")
    time_weekly.add_argument("--horizon-days", type=int, default=7)
    time_weekly.add_argument("--limit", type=int, default=20)
    time_weekly.add_argument("--write", action="store_true")
    time_weekly.set_defaults(func=cmd_time_weekly)
    time_agent_plan = time_sub.add_parser("agent-plan", help="Extract explicit Agent Tasks candidates from a time report.")
    time_agent_plan.add_argument("--from-report", required=True)
    time_agent_plan.add_argument("--write", action="store_true")
    time_agent_plan.set_defaults(func=cmd_time_agent_plan)
    time_agent_apply = time_sub.add_parser("agent-apply", help="Create Things Agent Tasks from a time action plan.")
    time_agent_apply.add_argument("--plan", required=True)
    time_agent_apply.set_defaults(func=cmd_time_agent_apply)
    time_verify = time_sub.add_parser("verify", help="Verify time planning prerequisites and state.")
    time_verify.set_defaults(func=cmd_time_verify)

    garden = sub.add_parser("garden", help="Run Personal OS gardening passes.")
    garden_sub = garden.add_subparsers(dest="garden_command", required=True)
    garden_daily = garden_sub.add_parser("daily", help="Run the daily Personal OS garden.")
    garden_daily.add_argument("--date", default="today")
    garden_daily.add_argument("--write", action="store_true")
    garden_daily.set_defaults(func=cmd_garden_daily)

    butler = sub.add_parser("butler", help="Rebuild Butler's Book views.")
    butler_sub = butler.add_subparsers(dest="butler_command", required=True)
    butler_rebuild = butler_sub.add_parser("rebuild", help="Rebuild the Butler's Book and related file views.")
    butler_rebuild.add_argument("--write", action="store_true")
    butler_rebuild.set_defaults(func=cmd_butler_rebuild)

    automation = sub.add_parser("automation", help="Install, enable, disable, or uninstall daily automation.")
    automation_sub = automation.add_subparsers(dest="automation_command", required=True)
    automation_install = automation_sub.add_parser("install", help="Install disabled LaunchAgent plist.")
    automation_install.set_defaults(func=cmd_automation_install)
    automation_enable = automation_sub.add_parser("enable", help="Enable and load the LaunchAgent.")
    automation_enable.set_defaults(func=cmd_automation_enable)
    automation_disable = automation_sub.add_parser("disable", help="Disable and unload the LaunchAgent.")
    automation_disable.set_defaults(func=cmd_automation_disable)
    automation_uninstall = automation_sub.add_parser("uninstall", help="Unload and remove the LaunchAgent plist.")
    automation_uninstall.set_defaults(func=cmd_automation_uninstall)

    scan = sub.add_parser("scan", help="Inspect Personal OS.")
    scan_sub = scan.add_subparsers(dest="scan_command", required=True)
    recent = scan_sub.add_parser("recent", help="List recent notes.")
    recent.add_argument("--limit", type=int, default=30)
    recent.set_defaults(func=cmd_scan_recent)

    index = sub.add_parser("index", help="Print or write the Personal OS index.")
    index.add_argument("--write", action="store_true")
    index.set_defaults(func=cmd_index)
    return parser


def main() -> int:
    args = build_parser().parse_args()
    args.func(args)
    return 0


if __name__ == "__main__":
    sys.exit(main())
